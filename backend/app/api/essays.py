import os
import json
import re
import tempfile
import zipfile
import uuid
import threading
from datetime import datetime, timedelta
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from sqlalchemy import func, or_, and_
from sqlalchemy.exc import IntegrityError

from ..database import get_db
from ..models.models import User, Essay, Course, EssayTask, OperationLog, SystemConfig, EssayImage
from ..schemas.schemas import EssayCreate, EssayOut, TaskOut, OperationLogOut
from ..utils.auth import get_current_user
from ..utils.file_utils import (
    get_essay_dir, generate_correction_filename,
    has_correction, count_corrections_in_dir, get_upload_dir, resize_image_within,
    safe_component,
)
from ..utils.ocr_utils import ocr_essay_images_with_fallback, ai_correct_text, ai_rewrite_text, count_cjk_chars
from ..utils.crypto_utils import load_config_row_value

router = APIRouter(prefix="/api/essays", tags=["作文"])


def _count_non_ws(text: str) -> int:
    """字数统计口径：不含空格/换行等空白字符，标点符号计入。"""
    return len([c for c in (text or "") if not c.isspace()])


def _essay_owned_filenames(essay, db=None):
    """该作文在磁盘上拥有的文件名（content_file + 图片），避免整目录误删/误搬其它作文的文件。"""
    names = set()
    if essay.content_file:
        names.add(os.path.basename(essay.content_file))
    if db is not None:
        imgs = db.query(EssayImage).filter(EssayImage.essay_id == essay.id).all()
        for img in imgs:
            if img.filename:
                names.add(img.filename)
    return names


def _delete_essay_disk_files(essay, db):
    """删除该作文自己的磁盘文件（原文 + 图片 + 对应修改文件），不动同目录下其它作文。"""
    if not essay.content_file:
        return
    base_dir = os.path.dirname(os.path.join(get_upload_dir(), essay.content_file))
    base = os.path.basename(essay.content_file)
    to_delete = _essay_owned_filenames(essay, db)
    to_delete.add(f"改_{base}")
    for fname in to_delete:
        p = os.path.join(base_dir, fname)
        try:
            if os.path.isfile(p):
                os.remove(p)
        except OSError:
            pass
    from ..utils.file_utils import _cleanup_empty_dirs
    _cleanup_empty_dirs(base_dir)


def _char_count_sql(col):
    """与 _count_non_ws 对应的 SQL 表达式（用于筛选/排序）。"""
    return func.char_length(func.regexp_replace(col, r"\s", "", "g"))


def _parse_uploaded_text(text: str):
    """解析上传的 docx/txt 文本：
    - 含「修改前/修改后」关键字 → 分别存入修改前/修改后（标题冒号可选，行首识别）
    - 不含 → 全部作为修改前
    返回 (content_text, corrected_text)
    """
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    if "修改后" not in text:
        m = re.search(r"(?m)^[ \t]*修改前(?:[：:][ \t]*|[ \t]+|[ \t]*(?=$))", text)
        if m:
            return text[m.end():].strip(), ""
        return text.strip(), ""

    def _heading_span(text_, keyword):
        m = re.search(r"(?m)^[ \t]*%s(?:[：:][ \t]*|[ \t]+|[ \t]*(?=$))" % keyword, text_)
        if m:
            return m.start(), m.end()
        m = re.search(r"%s[：:]\s*" % keyword, text_)
        return (m.start(), m.end()) if m else None

    after_span = _heading_span(text, "修改后")
    after = text[after_span[1]:].strip() if after_span else ""

    before_span = _heading_span(text, "修改前")
    if before_span:
        rest = text[before_span[1]:]
        after_in_rest = _heading_span(rest, "修改后")
        before = rest if after_in_rest is None else rest[:after_in_rest[0]]
        before = before.strip()
    else:
        before = text.split("修改后", 1)[0].strip()
    return before, after


def _parse_collect_time(s: str):
    """解析前端传入的收集时间（即作文列表「收集时间」列，覆盖 created_at；空/非法返回 None 保持默认）"""
    if not s:
        return None
    s = str(s).strip().replace("Z", "+00:00")
    try:
        parsed = datetime.fromisoformat(s)
    except ValueError:
        parsed = None
    if parsed is None:
        for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M", "%Y-%m-%d"):
            try:
                parsed = datetime.strptime(s, fmt)
                break
            except ValueError:
                continue
    if parsed is None:
        return None
    if parsed.tzinfo is not None:
        parsed = parsed.replace(tzinfo=None)
    return parsed


def _derive_title(text: str) -> str:
    """从作文文本自动分析标题：跳过「修改前/修改后」标题行与「——姓名」行，取首个有效行"""
    for line in (text or "").split("\n"):
        line = line.strip()
        if not line:
            continue
        if re.match(r"^修改[前后]\s*[：:]?", line):
            continue
        if line.startswith("——"):
            continue
        return line[:200]
    return ""


@router.get("/tasks", response_model=list[TaskOut])
def list_tasks(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """获取所有任务列表（供上传选择用）"""
    tasks = db.query(EssayTask).filter(EssayTask.deleted_at == None).order_by(EssayTask.created_at.desc()).all()
    return [TaskOut.model_validate(t) for t in tasks]


@router.get("/tasks/active", response_model=list[TaskOut])
def get_active_tasks(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """获取所有未过期的活跃收集任务（需登录）"""
    now = datetime.now()
    tasks = db.query(EssayTask).filter(
        EssayTask.is_active == True,
        EssayTask.deleted_at == None,
        (EssayTask.deadline == None) | (EssayTask.deadline >= now),
        (EssayTask.start_time == None) | (EssayTask.start_time <= now)
    ).all()
    return [TaskOut.model_validate(t) for t in tasks]


@router.get("/tasks/{task_id}/stats")
def get_task_stats(
    task_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """获取指定任务的统计数据"""
    task = db.query(EssayTask).filter(EssayTask.id == task_id).first()
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    
    total = db.query(func.count(Essay.id)).filter(
        Essay.task_id == task_id,
        Essay.deleted_at == None
    ).scalar() or 0
    pending = db.query(func.count(Essay.id)).filter(
        Essay.task_id == task_id,
        Essay.deleted_at == None,
        Essay.status.in_(["pending", "confirming", "rework"])
    ).scalar() or 0
    confirming = db.query(func.count(Essay.id)).filter(
        Essay.task_id == task_id,
        Essay.deleted_at == None,
        Essay.status == "confirming"
    ).scalar() or 0
    corrected = db.query(func.count(Essay.id)).filter(
        Essay.task_id == task_id,
        Essay.deleted_at == None,
        Essay.status == "corrected"
    ).scalar() or 0
    rework = db.query(func.count(Essay.id)).filter(
        Essay.task_id == task_id,
        Essay.deleted_at == None,
        Essay.status == "rework"
    ).scalar() or 0

    return {
        "task_id": task_id,
        "total": total,
        "pending": pending,
        "confirming": confirming,
        "corrected": corrected,
        "rework": rework
    }


@router.post("/tasks/stats")
def batch_task_stats(
    ids: list[int],
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """批量获取多个任务的统计数据（工作台用，避免逐个请求）"""
    result = []
    for task_id in ids:
        total = db.query(func.count(Essay.id)).filter(
            Essay.task_id == task_id,
            Essay.deleted_at == None
        ).scalar() or 0
        pending = db.query(func.count(Essay.id)).filter(
            Essay.task_id == task_id,
            Essay.deleted_at == None,
            Essay.status.in_(["pending", "confirming", "rework"])
        ).scalar() or 0
        confirming = db.query(func.count(Essay.id)).filter(
            Essay.task_id == task_id,
            Essay.deleted_at == None,
            Essay.status == "confirming"
        ).scalar() or 0
        corrected = db.query(func.count(Essay.id)).filter(
            Essay.task_id == task_id,
            Essay.deleted_at == None,
            Essay.status == "corrected"
        ).scalar() or 0
        rework = db.query(func.count(Essay.id)).filter(
            Essay.task_id == task_id,
            Essay.deleted_at == None,
            Essay.status == "rework"
        ).scalar() or 0
        result.append({
            "task_id": task_id,
            "total": total,
            "pending": pending,
            "confirming": confirming,
            "corrected": corrected,
            "rework": rework,
        })
    return result


@router.get("/existing-students")
def existing_students(
    task_id: int,
    essay_number: int = None,
    is_supplement: bool = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """批量上传预检：返回指定任务下已存在的学生姓名列表（用于跳过重复上传）"""
    q = db.query(Essay.student_name).filter(
        Essay.task_id == task_id,
        Essay.deleted_at == None,
    )
    if essay_number is not None:
        q = q.filter(Essay.essay_number == essay_number)
    if is_supplement is not None:
        q = q.filter(Essay.is_supplement == is_supplement)
    names = sorted({r[0] for r in q.all() if r[0]})
    return {"students": names}


def _build_download_filename(essay: Essay) -> str:
    """构建规范的下载文件名：标题——学生姓名年级第N次线上/线下补交（第几次为0或空时省略）"""
    title = essay.essay_title or "无标题"
    student = essay.student_name or "未知"
    grade = essay.grade or ""
    mode = essay.teaching_mode or "线下"
    supp = "补交" if essay.is_supplement else ""
    if essay.essay_number:
        return f"{title}——{student}{grade}第{essay.essay_number}次{mode}{supp}"
    return f"{title}——{student}{grade}{mode}{supp}"


def _generate_docx(essay: Essay, show_corrected: bool = False) -> str:
    """从 DB 生成 docx，返回临时文件路径。show_corrected=True 时包含修改后内容。"""
    from docx import Document

    doc = Document()
    _append_essay_to_doc(doc, essay, show_corrected=show_corrected)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_path = tmp.name
    tmp.close()
    doc.save(tmp_path)
    return tmp_path


def _append_essay_to_doc(doc, essay: Essay, show_corrected: bool = False, add_heading: bool = False, corrected_only: bool = False, original_only: bool = False) -> None:
    """把一篇作文的修改前后内容写入已有的 docx 文档。add_heading=True 时先加学生+标题行；corrected_only=True 时仅输出修改后内容；original_only=True 时仅输出修改前内容。"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml.ns import qn

    content = (essay.content_text or "").replace('\r\n', '\n').replace('\r', '\n')
    corrected = (essay.corrected_text or "").replace('\r\n', '\n').replace('\r', '\n')

    def _set_run_font(run):
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _set_para_format(para, is_title=False):
        fmt = para.paragraph_format
        fmt.line_spacing = Pt(12)
        fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        if not is_title:
            fmt.first_line_indent = Cm(0.74)
        else:
            fmt.first_line_indent = Cm(0)
            fmt.alignment = 1  # CENTER

    def _add_block(text, label):
        h = doc.add_paragraph()
        h_run = h.add_run(label)
        _set_run_font(h_run)
        _set_para_format(h, is_title=False)

        if not text.strip():
            return
        lines = [l.strip() for l in text.split('\n')]
        non_empty = [l for l in lines if l]
        for idx, line_text in enumerate(non_empty):
            p = doc.add_paragraph()
            run = p.add_run(line_text)
            _set_run_font(run)
            if idx < 2:
                run.bold = True
                _set_para_format(p, is_title=True)
            else:
                _set_para_format(p, is_title=False)

    if add_heading:
        head_text = f"{essay.student_name or '未知'}《{essay.essay_title or '无标题'}》"
        hp = doc.add_paragraph()
        hr = hp.add_run(head_text)
        _set_run_font(hr)
        hr.bold = True
        _set_para_format(hp, is_title=True)

    if corrected_only:
        _add_block(corrected, "修改后：")
        return

    if original_only:
        _add_block(content, "修改前：")
        return

    # 修改前
    _add_block(content, "修改前：")

    if show_corrected:
        # 分页符
        doc.add_page_break()
        # 修改后
        _add_block(corrected, "修改后：")


@router.get("/courses")
def list_courses_public(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """公开课程列表（收集者选课程用）"""
    courses = db.query(Course).filter(Course.deleted_at == None).all()
    return [{"id": c.id, "name": c.name} for c in courses]


@router.get("/collectors")
def list_collectors(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """获取收集者列表（用于筛选下拉）"""
    collectors = db.query(User).filter(
        User.deleted_at == None,
        (User.role.like('%collector%') | User.role.like('%admin%'))
    ).all()
    return [{"id": u.id, "nickname": u.nickname or u.username} for u in collectors]


@router.get("/recent-titles")
def recent_titles(
    limit: int = 5,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """获取最近上传的作文标题列表"""
    rows = (
        db.query(Essay.essay_title)
        .filter(Essay.deleted_at == None, Essay.essay_title != None, Essay.essay_title != "")
        .order_by(Essay.created_at.desc())
        .limit(limit * 3)
        .all()
    )
    seen = set()
    result = []
    for (title,) in rows:
        if title not in seen:
            seen.add(title)
            result.append(title)
        if len(result) >= limit:
            break
    return result


@router.get("/student-names")
def student_names(
    keyword: str = "",
    limit: int = 0,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """上传作文时学生姓名查找：返回已有学生姓名列表（limit=0 返回全部，供前端本地过滤）"""
    q = db.query(Essay.student_name).filter(
        Essay.deleted_at == None,
        Essay.student_name != None,
        Essay.student_name != "",
    )
    if keyword:
        q = q.filter(Essay.student_name.like(f"%{keyword}%"))
    names = sorted({r[0] for r in q.all() if r[0]})
    max_return = limit if limit and limit > 0 else 500
    names = names[:max_return]
    return {"names": names}


@router.get("/reviewers")
def list_reviewers(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """获取修改者列表（用于筛选下拉）"""
    reviewers = db.query(User).filter(
        User.deleted_at == None,
        (User.role.like('%reviewer%') | User.role.like('%admin%'))
    ).all()
    return [{"id": u.id, "nickname": u.nickname or u.username} for u in reviewers]


def _log_operation(db: Session, essay_id: int, user_id: int, action: str, detail: str = "",
                   old_value: str = "", new_value: str = "", batch_id: str = "", essay_ids: str = ""):
    try:
        log = OperationLog(
            essay_id=essay_id,
            user_id=user_id,
            action=action,
            detail=detail,
            old_value=old_value,
            new_value=new_value,
            batch_id=batch_id or None,
            essay_ids=essay_ids or None,
        )
        db.add(log)
    except Exception:
        pass


def build_file_path(db: Session, essay_data: dict) -> tuple[str, str, str, str]:
    """构建文件路径，返回 (dir_path, filename, year, month)"""
    now = datetime.now()
    year = str(now.year)
    month = f"{now.month}月"
    day = f"{now.day}"
    grade = essay_data.get("grade", "") or "未定年级"
    student_name = essay_data.get("student_name", "未知")

    dir_path = get_essay_dir(year, month, day, grade,
                              essay_data["essay_number"], student_name)
    os.makedirs(dir_path, exist_ok=True)
    return dir_path, year, month


@router.post("/upload", response_model=EssayOut)
async def upload_essay(
    essay_id: int = Form(None),
    task_id: int = Form(None),
    course_id: int = Form(None),
    grade: str = Form(""),
    essay_number: int = Form(1),
    essay_title: str = Form(""),
    student_name: str = Form(...),
    is_supplement: bool = Form(False),
    teaching_mode: str = Form("线下"),
    remark: str = Form(""),
    collector_note: str = Form(""),
    content_text: str = Form(""),
    collect_time: str = Form(None),
    mark_corrected: bool = Form(False),
    collected_by: int = Form(None),
    files: list[UploadFile] = File(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # 检查权限（暂时放宽：收集者直接通过）
    if "collector" not in current_user.role and "admin" not in current_user.role:
        raise HTTPException(status_code=403, detail="无权限")

    # 校验文件类型：.doc 旧版格式不支持（前端已拦截，后端兜底防止空作文落库）
    if files:
        for f in files:
            if f.filename and os.path.splitext(f.filename)[1].lower() == ".doc":
                raise HTTPException(status_code=400, detail="不支持 .doc 旧版格式，请另存为 .docx 后重新上传")

    # 确定收集者：管理员可指定，否则用当前用户
    collector_id = current_user.id
    if collected_by and "admin" in current_user.role:
        collector_id = collected_by

    # 确定课程：优先用传入 course_id，否则从任务继承
    effective_course_id = course_id
    if not effective_course_id and task_id:
        task = db.query(EssayTask).filter(EssayTask.id == task_id, EssayTask.deleted_at == None).first()
        if task:
            effective_course_id = task.course_id

    # 确定文件类型
    file_type = "text"
    content_file = ""

    # 创建或更新数据库记录
    if essay_id:
        essay = db.query(Essay).filter(Essay.id == essay_id).first()
        if not essay:
            raise HTTPException(status_code=404, detail="作文不存在")
        if "admin" not in current_user.role and essay.collected_by != current_user.id:
            raise HTTPException(status_code=403, detail="无权限编辑此作文")
        essay.task_id = task_id
        essay.course_id = effective_course_id
        essay.grade = grade
        essay.essay_number = essay_number
        essay.essay_title = essay_title
        essay.student_name = student_name
        essay.is_supplement = is_supplement
        essay.teaching_mode = teaching_mode
        essay.remark = remark
        essay.collector_note = collector_note
        if content_text:
            essay.content_text = content_text
        if mark_corrected:
            essay.status = "corrected"
            essay.corrected_at = datetime.now()
            essay.reviewer_id = current_user.id
        if essay.content_file and files:
            _delete_essay_disk_files(essay, db)
            db.query(EssayImage).filter(EssayImage.essay_id == essay.id).delete()
        try:
            db.commit()
        except IntegrityError:
            db.rollback()
            raise HTTPException(
                status_code=409,
                detail=f"学生「{student_name}」的同一篇作文已存在（第{essay_number}次{'补交' if is_supplement else '正篇'}），请先删除重复记录"
            )
        db.refresh(essay)
    else:
        try:
            essay = Essay(
                task_id=task_id,
                course_id=effective_course_id,
                grade=grade,
                essay_number=essay_number,
                essay_title=essay_title,
                student_name=student_name,
                is_supplement=is_supplement,
                teaching_mode=teaching_mode,
                remark=remark,
                collector_note=collector_note,
                content_text=content_text,
                file_type=file_type,
                collected_by=collector_id,
                status="corrected" if mark_corrected else "pending",
                corrected_at=datetime.now() if mark_corrected else None,
                reviewer_id=current_user.id if mark_corrected else None,
            )
            db.add(essay)
            db.flush()
        except IntegrityError:
            db.rollback()
            raise HTTPException(
                status_code=409,
                detail=f"学生「{student_name}」的同一篇作文已存在（第{essay_number}次{'补交' if is_supplement else '正篇'}），如需修改请先删除旧记录"
            )
    _log_operation(db, essay.id, current_user.id, "上传", student_name)
    db.commit()
    db.refresh(essay)

    # 保存文件
    now = datetime.now()
    ts = now.strftime("%H%M%S")

    grade_name = safe_component(grade, "未定年级")
    if teaching_mode:
        grade_name = f"{grade_name}{safe_component(teaching_mode, '')}"

    dir_path = os.path.join(
        get_upload_dir(),
        safe_component(str(now.year), "0000"),
        safe_component(f"{now.month}月", "1月"),
        safe_component(str(now.day), "1"),
        f"{grade_name}第{essay_number}次" if essay_number not in (None, 0) else grade_name,
        safe_component(student_name, "未知"),
    )
    title_component = safe_component(essay_title, "无标题")
    if is_supplement:
        title_component += "_补交"
    dir_path = os.path.join(dir_path, title_component)
    os.makedirs(dir_path, exist_ok=True)

    if files:
        img_index = 1
        uploaded_files = []
        text_buffer = []
        for f in files:
            if not f.filename:
                continue
            ext = os.path.splitext(f.filename)[1].lower()
            content = await f.read()
            if ext in [".jpg", ".jpeg", ".png", ".gif", ".webp"]:
                essay.file_type = "image"
                content = resize_image_within(content)
                img_name = f"{img_index}{ext}"
                img_index += 1
                img_path = os.path.join(dir_path, img_name)
                with open(img_path, "wb") as fw:
                    fw.write(content)
                uploaded_files.append(img_name)
                essay_image = EssayImage(essay_id=essay.id, filename=img_name, image_data=content)
                db.add(essay_image)
            elif ext == ".docx":
                # docx 不保存本地，解析文本（含 修改前/修改后 拆分）
                try:
                    from docx import Document
                    import io
                    doc = Document(io.BytesIO(content))
                    text_lines = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            text_lines.append(para.text.strip())
                    if text_lines:
                        text_buffer.append("\n".join(text_lines))
                except Exception:
                    pass
            elif ext == ".txt":
                # txt 直接读取文本（含 修改前/修改后 拆分）
                try:
                    text = content.decode("utf-8", errors="replace").strip()
                    if text:
                        text_buffer.append(text)
                except Exception:
                    pass

        if text_buffer:
            # 多个 docx/txt 逐个拆分后合并：含「修改后」的文件作为修改版（修改前/修改后分别合并），
            # 全部无修改标记时才把所有文本合并为修改前
            parsed_parts = [_parse_uploaded_text(t) for t in text_buffer]
            marked_parts = [(b, a) for b, a in parsed_parts if a]
            if marked_parts:
                before_text = "\n".join(b for b, a in marked_parts if b)
                after_text = "\n".join(a for b, a in marked_parts if a)
            else:
                before_text = "\n".join(b for b, a in parsed_parts if b)
                after_text = ""
            if before_text:
                essay.content_text = before_text
            if after_text:
                essay.corrected_text = after_text

        if uploaded_files:
            essay.content_file = os.path.relpath(
                os.path.join(dir_path, uploaded_files[0]), get_upload_dir()
            )

    parsed_ts = _parse_collect_time(collect_time)
    if parsed_ts:
        essay.created_at = parsed_ts
        if mark_corrected:
            essay.corrected_at = parsed_ts

    if not essay.essay_title and essay.content_text:
        derived = _derive_title(essay.content_text)
        if derived:
            essay.essay_title = derived

    db.commit()
    db.refresh(essay)
    return _essay_to_out(essay, db)


@router.post("/upload-correction-docx")
async def upload_correction_docx(
    grade: str = Form(...),
    essay_number: int = Form(...),
    teaching_mode: str = Form("线下"),
    student_name: str = Form(...),
    essay_title: str = Form(""),
    content_text: str = Form(""),
    corrected_text: str = Form(""),
    collect_time: str = Form(None),
    mark_corrected: bool = Form(False),
    is_supplement: bool = Form(False),
    task_id: int = Form(None),
    course_id: int = Form(None),
    collected_by: int = Form(None),
    collector_note: str = Form(""),
    file: UploadFile = File(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """批量上传修改后docx：保存文件到年级目录，创建作文记录"""
    if "collector" not in current_user.role and "admin" not in current_user.role:
        raise HTTPException(status_code=403, detail="无权限")

    # 确定收集者：管理员可指定，否则用当前用户
    collector_id = current_user.id
    if collected_by and "admin" in current_user.role:
        collector_id = collected_by

    # 确定课程：优先用传入 course_id，否则从任务继承
    effective_course_id = course_id
    if not effective_course_id and task_id:
        task = db.query(EssayTask).filter(EssayTask.id == task_id, EssayTask.deleted_at == None).first()
        if task:
            effective_course_id = task.course_id

    now = datetime.now()
    grade_name = safe_component(grade, "未定年级")
    if teaching_mode:
        grade_name = f"{grade_name}{safe_component(teaching_mode, '')}"

    dir_path = os.path.join(
        get_upload_dir(),
        safe_component(str(now.year), "0000"),
        safe_component(f"{now.month}月", "1月"),
        safe_component(str(now.day), "1"),
        f"{grade_name}第{essay_number}次" if essay_number not in (None, 0) else grade_name,
    )

    try:
        os.makedirs(dir_path, exist_ok=True)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"创建目录失败: {str(e)}")

    file_saved = False
    if file and file.filename:
        safe_filename = os.path.basename(file.filename)
        file_path = os.path.join(dir_path, safe_filename)
        try:
            content = await file.read()
            with open(file_path, "wb") as fw:
                fw.write(content)
            file_saved = True
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"保存文件失败: {str(e)}")

    try:
        # 检查是否已存在同一条记录（同一学生同一次作文，严格按任务隔离）
        existing_query = db.query(Essay).filter(
            Essay.student_name == student_name,
            Essay.essay_number == essay_number,
            Essay.is_supplement == is_supplement,
            Essay.deleted_at == None,
        )
        if task_id is not None:
            existing_query = existing_query.filter(Essay.task_id == task_id)
        else:
            existing_query = existing_query.filter(Essay.task_id.is_(None))
        existing = existing_query.first()

        if existing:
            # 更新已有记录
            existing.essay_title = essay_title or existing.essay_title
            existing.content_text = content_text or existing.content_text
            existing.corrected_text = corrected_text if corrected_text else existing.corrected_text
            if mark_corrected:
                existing.status = "corrected"
                existing.corrected_at = datetime.now()
                existing.reviewer_id = current_user.id
            else:
                existing.status = "confirming" if corrected_text and existing.status == "pending" and existing.content_text and existing.content_text.strip() else existing.status
                existing.corrected_at = datetime.now() if corrected_text else existing.corrected_at
                existing.reviewer_id = current_user.id if corrected_text else existing.reviewer_id
            existing.teaching_mode = teaching_mode or existing.teaching_mode
            if collected_by:
                existing.collected_by = collector_id
            existing.is_supplement = is_supplement
            if collector_note:
                existing.collector_note = collector_note
            if task_id is not None:
                existing.task_id = task_id
            if effective_course_id:
                existing.course_id = effective_course_id
            essay = existing
        else:
            # 新建记录
            essay = Essay(
                course_id=effective_course_id,
                grade=grade,
                essay_number=essay_number,
                essay_title=essay_title,
                student_name=student_name,
                is_supplement=is_supplement,
                teaching_mode=teaching_mode,
                remark="",
                collector_note=collector_note,
                content_text=content_text,
                corrected_text=corrected_text if corrected_text else "",
                file_type="docx",
                collected_by=collector_id,
                task_id=task_id,
                status="corrected" if mark_corrected else ("confirming" if corrected_text and content_text and content_text.strip() else "pending"),
                corrected_at=datetime.now() if (mark_corrected or corrected_text) else None,
                reviewer_id=current_user.id if (mark_corrected or corrected_text) else None,
            )
            db.add(essay)

        db.flush()
        parsed_ts = _parse_collect_time(collect_time)
        if parsed_ts:
            essay.created_at = parsed_ts
            if mark_corrected:
                essay.corrected_at = parsed_ts
        if not essay.essay_title:
            essay.essay_title = _derive_title(essay.content_text) or _derive_title(essay.corrected_text)
        _log_operation(db, essay.id, current_user.id, "修改", student_name)
        db.commit()
        db.refresh(essay)
    except IntegrityError:
        db.rollback()
        raise HTTPException(
            status_code=409,
            detail=f"学生「{student_name}」第{essay_number}次作文在所选任务下已存在，请先删除重复记录"
        )
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"创建记录失败: {str(e)}")

    return {"message": "上传成功", "id": essay.id}


@router.get("")
def list_essays(
    status: str = None,
    name: str = None,
    grade: str = None,
    essay_number: int = None,
    teaching_mode: str = None,
    collected_by: int = None,
    course_id: int = None,
    remark: str = None,
    essay_title: str = None,
    task_id: int = None,
    reviewer_id: int = None,
    is_supplement: bool = None,
    task_name: str = None,
    word_count_min: int = None,
    word_count_max: int = None,
    corrected_word_count_min: int = None,
    corrected_word_count_max: int = None,
    date_from: str = None,
    date_to: str = None,
    corrected_from: str = None,
    corrected_to: str = None,
    sort_by: str = "created_at",
    sort_order: str = "desc",
    page: int = 1,
    page_size: int = 50,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    q = db.query(Essay).filter(Essay.deleted_at == None)

    # 权限过滤：收集者和游客可以查看所有作文，批改者只能看自己的
    if "admin" not in current_user.role and "guest" not in current_user.role:
        if "reviewer" in current_user.role and "collector" not in current_user.role:
            q = q.filter(Essay.reviewer_id == current_user.id)
        # 收集者可以查看所有作文，不做过滤

    if course_id:
        q = q.filter(Essay.course_id == course_id)
    if name:
        q = q.filter(Essay.student_name.like(f"%{name}%"))
    if grade:
        q = q.filter(Essay.grade == grade)
    if essay_number:
        q = q.filter(Essay.essay_number == essay_number)
    if teaching_mode:
        q = q.filter(Essay.teaching_mode == teaching_mode)
    if remark:
        q = q.filter(Essay.collector_note.like(f"%{remark}%"))
    if collected_by:
        q = q.filter(Essay.collected_by == collected_by)
    if essay_title:
        q = q.filter(Essay.essay_title.like(f"%{essay_title}%"))
    if task_id == 0:
        q = q.filter((Essay.task_id.is_(None)) | (Essay.task_id == 0))
    elif task_id is not None:
        q = q.filter(Essay.task_id == task_id)
    if reviewer_id is not None:
        q = q.filter(Essay.reviewer_id == reviewer_id)
    if is_supplement is not None:
        q = q.filter(Essay.is_supplement == is_supplement)
    if task_name:
        q = q.join(EssayTask, Essay.task_id == EssayTask.id, isouter=True).filter(EssayTask.name.like(f"%{task_name}%"))
    if word_count_min is not None:
        q = q.filter(_char_count_sql(Essay.content_text) >= word_count_min)
    if word_count_max is not None:
        q = q.filter(_char_count_sql(Essay.content_text) <= word_count_max)
    if corrected_word_count_min is not None:
        q = q.filter(_char_count_sql(Essay.corrected_text) >= corrected_word_count_min)
    if corrected_word_count_max is not None:
        q = q.filter(_char_count_sql(Essay.corrected_text) <= corrected_word_count_max)
    if date_from:
        q = q.filter(Essay.created_at >= date_from)
    if date_to:
        q = q.filter(Essay.created_at <= date_to + " 23:59:59")
    if corrected_from:
        q = q.filter(Essay.corrected_at >= corrected_from)
    if corrected_to:
        q = q.filter(Essay.corrected_at <= corrected_to + " 23:59:59")

    # 排序
    from sqlalchemy import case
    allowed_sort = {"created_at": Essay.created_at, "corrected_at": Essay.corrected_at, "student_name": Essay.student_name, "grade": Essay.grade, "essay_number": Essay.essay_number, "status": Essay.status, "remark": Essay.remark, "is_supplement": Essay.is_supplement}
    
    # 处理特殊排序字段
    if sort_by == "collector_name":
        q = q.outerjoin(User, User.id == Essay.collected_by)
        order_col = User.nickname
    elif sort_by == "reviewer_name":
        q = q.outerjoin(User, User.id == Essay.reviewer_id)
        order_col = User.nickname
    elif sort_by == "word_count":
        order_col = _char_count_sql(Essay.content_text)
    elif sort_by == "corrected_word_count":
        order_col = _char_count_sql(Essay.corrected_text)
    else:
        order_col = allowed_sort.get(sort_by, Essay.created_at)
    
    if sort_order == "asc":
        q = q.order_by(order_col.asc(), Essay.id.asc())
    else:
        q = q.order_by(order_col.desc(), Essay.id.desc())

    # 统计随当前筛选计算（不含状态筛选本身）
    pending_total = q.filter(Essay.status.in_(["pending", "confirming", "rework"])).count()
    corrected_total = q.filter(Essay.status == "corrected").count()
    if status:
        q = q.filter(Essay.status == status)
    total = q.count()
    q = q.offset((page - 1) * page_size).limit(page_size)
    essays = q.all()
    result = _essay_to_out_bulk(essays, db)

    # 收集者列表（用于前端下拉筛选）
    collectors = db.query(User).filter(
        User.deleted_at == None,
        (User.role.like('%collector%') | User.role.like('%admin%'))
    ).all()
    collector_list = [{"id": u.id, "nickname": u.nickname or u.username} for u in collectors]

    return {
        "items": result,
        "total": total,
        "pending": pending_total,
        "corrected": corrected_total,
        "collectors": collector_list,
        "page": page,
        "page_size": page_size,
    }


@router.get("/pending")
def pending_essays(
    name: str = None,
    grade: str = None,
    essay_number: int = None,
    teaching_mode: str = None,
    task_id: int = None,
    task_name: str = None,
    collected_by: int = None,
    essay_title: str = None,
    status: str = None,
    date_from: str = None,
    date_to: str = None,
    sort_by: str = "created_at",
    sort_order: str = "asc",
    page: int = 1,
    page_size: int = 50,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """获取待修改的作文列表（修改者/游客用），支持筛选和分页"""
    if "reviewer" not in current_user.role and "admin" not in current_user.role and "guest" not in current_user.role:
        raise HTTPException(status_code=403, detail="无权限")

    q = db.query(Essay).filter(
        Essay.deleted_at == None,
    )

    if status:
        q = q.filter(Essay.status == status)
    else:
        q = q.filter(Essay.status.in_(["pending", "confirming", "rework"]))
    if name:
        q = q.filter(Essay.student_name.like(f"%{name}%"))
    if grade:
        q = q.filter(Essay.grade == grade)
    if essay_number:
        q = q.filter(Essay.essay_number == essay_number)
    if teaching_mode:
        q = q.filter(Essay.teaching_mode == teaching_mode)
    if task_id is not None:
        if task_id == 0:
            q = q.filter((Essay.task_id.is_(None)) | (Essay.task_id == 0))
        else:
            q = q.filter(Essay.task_id == task_id)
    if task_name:
        q = q.join(EssayTask, Essay.task_id == EssayTask.id, isouter=True).filter(EssayTask.name.like(f"%{task_name}%"))
    if collected_by:
        q = q.filter(Essay.collected_by == collected_by)
    if essay_title:
        q = q.filter(Essay.essay_title.like(f"%{essay_title}%"))
    if date_from:
        q = q.filter(Essay.created_at >= date_from)
    if date_to:
        q = q.filter(Essay.created_at <= date_to + " 23:59:59")

    allowed_sort = {
        "created_at": Essay.created_at,
        "student_name": Essay.student_name,
        "grade": Essay.grade,
        "essay_number": Essay.essay_number,
        "teaching_mode": Essay.teaching_mode,
        "word_count": _char_count_sql(Essay.content_text),
        "corrected_word_count": _char_count_sql(Essay.corrected_text),
    }
    order_col = allowed_sort.get(sort_by, Essay.created_at)
    if sort_order == "desc":
        q = q.order_by(order_col.desc())
    else:
        q = q.order_by(order_col.asc())

    total = q.count()
    essays = q.offset((page - 1) * page_size).limit(page_size).all()
    result = _essay_to_out_bulk(essays, db)
    return {"items": result, "total": total, "page": page, "page_size": page_size}


@router.get("/pending/next")
def next_pending_essay(
    current_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """获取当前作文之后的下一篇未修改作文（批改流水线跳转用）"""
    if "reviewer" not in current_user.role and "admin" not in current_user.role and "guest" not in current_user.role:
        raise HTTPException(status_code=403, detail="无权限")
    cur = db.query(Essay).filter(Essay.id == current_id).first()
    if not cur:
        return {"next_id": None}
    cur_created = cur.created_at or datetime.min
    next_row = (
        db.query(Essay.id)
        .filter(
            Essay.deleted_at == None,
            Essay.status.in_(["pending", "confirming", "rework"]),
            or_(
                Essay.created_at > cur_created,
                and_(Essay.created_at == cur_created, Essay.id > cur.id),
            ),
        )
        .order_by(Essay.created_at.asc(), Essay.id.asc())
        .first()
    )
    return {"next_id": next_row[0] if next_row else None}


@router.get("/trash")
def list_trash(
    page: int = 1,
    page_size: int = 50,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """查看已删除的作文（管理员）"""
    if "admin" not in current_user.role:
        raise HTTPException(status_code=403, detail="无权限")
    q = db.query(Essay).filter(Essay.deleted_at != None).order_by(Essay.deleted_at.desc())
    total = q.count()
    essays = q.offset((page - 1) * page_size).limit(page_size).all()
    return {
        "items": _essay_to_out_bulk(essays, db),
        "total": total,
        "page": page,
        "page_size": page_size,
    }


@router.post("/{essay_id}/restore")
def restore_essay(
    essay_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """从回收站恢复作文"""
    if "admin" not in current_user.role:
        raise HTTPException(status_code=403, detail="无权限")
    essay = db.query(Essay).filter(Essay.id == essay_id, Essay.deleted_at != None).first()
    if not essay:
        raise HTTPException(status_code=404, detail="作文不存在或不在回收站")
    essay.deleted_at = None
    _log_operation(db, essay.id, current_user.id, "恢复", essay.student_name)
    db.commit()
    return {"message": "已恢复"}


@router.get("/stats")
def essay_stats(
    year: int = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Dashboard 统计数据（year 指定热力图年份，缺省为近365天）"""
    now = datetime.now()
    today = now.date()
    month_start = today.replace(day=1)

    base = db.query(Essay).filter(Essay.deleted_at == None)
    total = base.count()
    pending = base.filter(Essay.status.in_(["pending", "confirming", "rework"])).count()
    confirming = base.filter(Essay.status == "confirming").count()
    rework = base.filter(Essay.status == "rework").count()
    corrected = base.filter(Essay.status == "corrected").count()
    this_month = base.filter(Essay.created_at >= month_start).count()

    grade_rows = (
        base.with_entities(Essay.grade, func.count(Essay.id))
        .group_by(Essay.grade)
        .order_by(func.count(Essay.id).desc())
        .all()
    )
    grade_dist = [{"name": g or "未知", "value": c} for g, c in grade_rows]

    class_rows = (
        base.with_entities(Course.name, func.count(Essay.id))
        .join(Course, Course.id == Essay.course_id)
        .filter(~Course.name.like("%迁移%"))
        .group_by(Course.id, Course.name)
        .order_by(func.count(Essay.id).desc())
        .all()
    )
    class_dist = [{"name": n, "value": c} for n, c in class_rows]

    collector_rows = (
        base.with_entities(User.nickname, User.username, func.count(Essay.id))
        .join(User, User.id == Essay.collected_by)
        .filter(User.role.like("%collector%"))
        .group_by(Essay.collected_by, User.nickname, User.username)
        .order_by(func.count(Essay.id).desc())
        .limit(10)
        .all()
    )
    collector_rank = [{"name": n or u, "value": c} for n, u, c in collector_rows]

    # 近14天有上传的收集者（用于趋势图按人分组展示）
    window_start = datetime.combine(today - timedelta(days=13), datetime.min.time())
    collector_rows = (
        db.query(Essay.collected_by, User.nickname, User.username, func.count(Essay.id).label("cnt"))
        .join(User, User.id == Essay.collected_by)
        .filter(Essay.deleted_at == None, Essay.created_at >= window_start)
        .group_by(Essay.collected_by, User.nickname, User.username)
        .order_by(func.count(Essay.id).desc())
        .limit(8)
        .all()
    )
    trend_collectors = [{"id": cid, "name": nickname or username} for cid, nickname, username, _ in collector_rows]
    trend_collector_ids = [cid for cid, _, _, _ in collector_rows]

    day_collector_rows = (
        db.query(func.date(Essay.created_at).label("day"), Essay.collected_by, func.count(Essay.id).label("cnt"))
        .filter(
            Essay.deleted_at == None,
            Essay.created_at >= window_start,
            Essay.collected_by.in_(trend_collector_ids or [0]),
        )
        .group_by("day", Essay.collected_by)
        .all()
    )
    day_collector_map = {}
    for day, cid, cnt in day_collector_rows:
        day_collector_map.setdefault(str(day), {})[cid] = cnt

    trend = []
    for i in range(13, -1, -1):
        d = today - timedelta(days=i)
        d_start = datetime.combine(d, datetime.min.time())
        d_end = datetime.combine(d, datetime.max.time())
        uploaded = base.filter(Essay.created_at >= d_start, Essay.created_at <= d_end).count()
        done = base.filter(Essay.corrected_at >= d_start, Essay.corrected_at <= d_end).count()
        by_collector = {str(cid): day_collector_map.get(d.strftime("%Y-%m-%d"), {}).get(cid, 0) for cid in trend_collector_ids}
        trend.append({
            "date": d.strftime("%m-%d"),
            "uploaded": uploaded,
            "corrected": done,
            "by_collector": by_collector,
        })

    # GitHub 风格：按年每日上传数量（贡献热力图）
    # 可切换年份；缺省展示近365天
    if year:
        year_start = datetime(year, 1, 1)
        year_end = datetime(year, 12, 31, 23, 59, 59)
        daily_rows = (
            db.query(func.date(Essay.created_at).label("day"), func.count(Essay.id).label("cnt"))
            .filter(Essay.deleted_at == None, Essay.created_at >= year_start, Essay.created_at <= year_end)
            .group_by("day")
            .all()
        )
        daily_map = {str(day): cnt for day, cnt in daily_rows}
        daily_upload = []
        is_leap = (year % 4 == 0 and year % 100 != 0) or (year % 400 == 0)
        total_days = 366 if is_leap else 365
        for i in range(total_days):
            d = (year_start + timedelta(days=i)).date()
            daily_upload.append({"date": d.strftime("%Y-%m-%d"), "count": daily_map.get(d.strftime("%Y-%m-%d"), 0)})
    else:
        year_start = datetime.combine(today - timedelta(days=364), datetime.min.time())
        daily_rows = (
            db.query(func.date(Essay.created_at).label("day"), func.count(Essay.id).label("cnt"))
            .filter(Essay.deleted_at == None, Essay.created_at >= year_start)
            .group_by("day")
            .all()
        )
        daily_map = {str(day): cnt for day, cnt in daily_rows}
        daily_upload = []
        for i in range(364, -1, -1):
            d = today - timedelta(days=i)
            daily_upload.append({"date": d.strftime("%Y-%m-%d"), "count": daily_map.get(d.strftime("%Y-%m-%d"), 0)})

    # 可用的年份列表（有上传记录的年份，用于热力图年份切换）
    year_rows = (
        db.query(func.extract("year", Essay.created_at).label("yr"))
        .filter(Essay.deleted_at == None)
        .distinct()
        .all()
    )
    available_years = sorted([int(yr) for yr, in year_rows if yr is not None], reverse=True)

    return {
        "total": total,
        "pending": pending,
        "confirming": confirming,
        "rework": rework,
        "corrected": corrected,
        "this_month": this_month,
        "grade_dist": grade_dist,
        "class_dist": class_dist,
        "collector_rank": collector_rank,
        "trend": trend,
        "trend_collectors": trend_collectors,
        "daily_upload": daily_upload,
        "available_years": available_years,
    }


@router.get("/my-stats")
def my_stats(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """当前用户的个人统计（个人信息页用，按角色区分收集/批改/总览视角）"""
    mine = db.query(Essay).filter(
        Essay.collected_by == current_user.id,
        Essay.deleted_at == None,
    )
    reviewed = db.query(Essay).filter(
        Essay.reviewer_id == current_user.id,
        Essay.deleted_at == None,
    )
    base = db.query(Essay).filter(Essay.deleted_at == None)
    today_start = datetime.combine(datetime.now().date(), datetime.min.time())
    return {
        # 收集视角
        "collected_total": mine.count(),
        "collected_pending": mine.filter(Essay.status == "pending").count(),
        "collected_confirming": mine.filter(Essay.status == "confirming").count(),
        "collected_rework": mine.filter(Essay.status == "rework").count(),
        "collected_corrected": mine.filter(Essay.status == "corrected").count(),
        "uploaded_today": mine.filter(Essay.created_at >= today_start).count(),
        # 批改视角
        "reviewed_total": reviewed.count(),
        "reviewed_confirming": reviewed.filter(Essay.status == "confirming").count(),
        "reviewed_rework": reviewed.filter(Essay.status == "rework").count(),
        "reviewed_corrected": reviewed.filter(Essay.status == "corrected").count(),
        "reviewed_today": reviewed.filter(Essay.corrected_at >= today_start).count(),
        # 系统总览
        "sys_total": base.count(),
        "sys_pending": base.filter(Essay.status == "pending").count(),
        "sys_confirming": base.filter(Essay.status == "confirming").count(),
        "sys_rework": base.filter(Essay.status == "rework").count(),
        "sys_corrected": base.filter(Essay.status == "corrected").count(),
        # 当前待处理总数（批改待办）
        "todo_total": base.filter(Essay.status.in_(["pending", "confirming", "rework"])).count(),
    }


@router.get("/download/by-course/{course_id}")
def download_by_course(
    course_id: int,
    essay_number: int = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """按课程打包下载全部作文"""
    if "guest" in current_user.role:
        raise HTTPException(status_code=403, detail="游客无下载权限")
    cls = db.query(Course).filter(Course.id == course_id).first()
    if not cls:
        raise HTTPException(status_code=404, detail="课程不存在")

    q = db.query(Essay).filter(Essay.course_id == course_id, Essay.deleted_at == None)
    if essay_number:
        q = q.filter(Essay.essay_number == essay_number)

    essays = q.all()
    if not essays:
        raise HTTPException(status_code=404, detail="没有找到作文")

    dirs = set()
    for e in essays:
        if e.content_file:
            d = os.path.dirname(os.path.join(get_upload_dir(), e.content_file))
            dirs.add(d)

    tmp_dir = tempfile.mkdtemp()
    archive_name = f"{cls.name}_作文打包"
    if essay_number:
        archive_name += f"_第{essay_number}次"
    archive_name += ".tar.gz"

    archive_path = os.path.join(tmp_dir, archive_name)

    import tarfile
    with tarfile.open(archive_path, "w:gz") as tar:
        for d in dirs:
            if os.path.exists(d):
                tar.add(d, arcname=os.path.relpath(d, get_upload_dir()))

    from starlette.background import BackgroundTask

    def _cleanup_course_archive(tmp_dir, archive_path):
        try:
            if os.path.exists(archive_path):
                os.remove(archive_path)
            os.rmdir(tmp_dir)
        except OSError:
            pass

    return FileResponse(archive_path, filename=archive_name, media_type="application/gzip",
                        background=BackgroundTask(_cleanup_course_archive, tmp_dir, archive_path))


# ===== 以下所有 /{essay_id}/xxx 具名路由必须在 /{essay_id} 通用路由之前 =====


@router.post("/{essay_id}/claim")
def claim_essay(
    essay_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """修改者认领作文"""
    if "reviewer" not in current_user.role and "admin" not in current_user.role:
        raise HTTPException(status_code=403, detail="无权限")

    essay = db.query(Essay).filter(Essay.id == essay_id).first()
    if not essay:
        raise HTTPException(status_code=404, detail="作文不存在")
    if essay.reviewer_id:
        raise HTTPException(status_code=400, detail="该作文已被其他人认领")

    essay.reviewer_id = current_user.id
    _log_operation(db, essay.id, current_user.id, "修改", essay.student_name)
    db.commit()
    return {"message": "认领成功"}


@router.delete("/{essay_id}")
def delete_essay(
    essay_id: int,
    delete_file: bool = False,
    permanent: bool = False,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """删除作文。
    - 默认：软删除（写入 deleted_at），文件保留
    - delete_file=true：同时删除本地文件（需管理员）
    - permanent=true：物理删除数据库记录（需管理员）
    """
    essay = db.query(Essay).filter(Essay.id == essay_id).first()
    if not essay:
        raise HTTPException(status_code=404, detail="作文不存在")
    if "admin" not in current_user.role and essay.collected_by != current_user.id:
        raise HTTPException(status_code=403, detail="无权限删除此作文")

    if delete_file:
        if "admin" not in current_user.role:
            raise HTTPException(status_code=403, detail="仅管理员可删除本地文件")
        if essay.content_file:
            _delete_essay_disk_files(essay, db)
        essay.content_file = ""

    if permanent:
        if "admin" not in current_user.role:
            raise HTTPException(status_code=403, detail="仅管理员可彻底删除")
        db.query(EssayImage).filter(EssayImage.essay_id == essay_id).delete()
        db.delete(essay)
        _log_operation(db, essay_id, current_user.id, "删除", essay.student_name)
        db.commit()
        return {"message": "已彻底删除"}

    essay.deleted_at = datetime.now()
    _log_operation(db, essay_id, current_user.id, "删除", essay.student_name)
    db.commit()
    return {"message": "已移入回收站"}


@router.post("/batch-delete")
def batch_delete_essays(
    payload: dict,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """批量删除作文。delete_file/permanent 语义与单条删除一致（均需管理员）。"""
    ids = payload.get("ids") or []
    delete_file = bool(payload.get("delete_file", False))
    permanent = bool(payload.get("permanent", False))
    if not ids or not isinstance(ids, list):
        raise HTTPException(status_code=400, detail="参数错误")

    done = 0
    errors = []
    for essay_id in ids:
        essay = db.query(Essay).filter(Essay.id == essay_id).first()
        if not essay:
            continue
        if "admin" not in current_user.role and essay.collected_by != current_user.id:
            errors.append({"id": essay_id, "detail": "无权限删除此作文"})
            continue
        if delete_file:
            if "admin" not in current_user.role:
                errors.append({"id": essay_id, "detail": "仅管理员可删除本地文件"})
                continue
            if essay.content_file:
                _delete_essay_disk_files(essay, db)
            essay.content_file = ""
        if permanent:
            if "admin" not in current_user.role:
                errors.append({"id": essay_id, "detail": "仅管理员可彻底删除"})
                continue
            db.query(EssayImage).filter(EssayImage.essay_id == essay_id).delete()
            db.delete(essay)
            _log_operation(db, essay_id, current_user.id, "删除", essay.student_name)
        else:
            essay.deleted_at = datetime.now()
            _log_operation(db, essay_id, current_user.id, "删除", essay.student_name)
        done += 1
    db.commit()
    return {"success": done, "errors": errors, "total": len(ids)}


@router.post("/{essay_id}/upload-correction")
async def upload_correction(
    essay_id: int,
    file: UploadFile = File(None),
    corrected_text: str = Form(""),
    reviewer_note: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """上传修改结果（支持文件上传 + 文字修改）"""
    if "reviewer" not in current_user.role and "admin" not in current_user.role:
        raise HTTPException(status_code=403, detail="无权限")
    essay = db.query(Essay).filter(Essay.id == essay_id).first()
    if not essay:
        raise HTTPException(status_code=404, detail="作文不存在")

    # 至少提供文件或文字
    if not file and not corrected_text.strip():
        raise HTTPException(status_code=400, detail="请上传文件或填写修改文字")

    # 保存文件（如果有）
    corr_name = ""
    if file and file.filename:
        if not essay.content_file:
            raise HTTPException(status_code=400, detail="原文不存在，无法上传修改")
        original_path = os.path.join(get_upload_dir(), essay.content_file)
        original_dir = os.path.dirname(original_path)
        original_name = os.path.basename(original_path)

        corr_name = generate_correction_filename(original_name)
        corr_path = os.path.join(original_dir, corr_name)

        content = await file.read()
        with open(corr_path, "wb") as f:
            f.write(content)

    # 保存文字修改（如果有）
    if corrected_text.strip():
        essay.corrected_text = corrected_text.strip()

    if reviewer_note.strip():
        essay.reviewer_note = reviewer_note.strip()

    essay.reviewer_id = current_user.id
    if essay.status in ("pending", "rework") and essay.content_text and essay.content_text.strip():
        essay.status = "confirming"
    essay.corrected_at = datetime.now()
    _log_operation(db, essay.id, current_user.id, "修改", essay.student_name)
    db.commit()

    return {"message": "修改上传成功", "file": corr_name, "corrected_text": essay.corrected_text}


@router.get("/{essay_id}/images")
def get_essay_images(
    essay_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """获取作文目录下的所有图片（返回URL列表）"""
    essay = db.query(Essay).filter(Essay.id == essay_id).first()
    if not essay:
        return {"images": []}

    images = set()

    if essay.content_file:
        dir_path = os.path.dirname(os.path.join(get_upload_dir(), essay.content_file))
        if os.path.exists(dir_path):
            for f in os.listdir(dir_path):
                if f.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.webp')):
                    images.add(f)

    db_images = db.query(EssayImage).filter(EssayImage.essay_id == essay_id).all()
    for img in db_images:
        images.add(img.filename)

    sorted_images = sorted(images)
    base_url = "/api/essays/" + str(essay_id) + "/file/"
    return {"images": [base_url + img for img in sorted_images], "dir": ""}


@router.get("/{essay_id}/file/{filename}")
def get_essay_file(
    essay_id: int,
    filename: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """获取作文目录下的单个文件（需登录，图片经前端转 blob 后展示）"""
    essay = db.query(Essay).filter(Essay.id == essay_id).first()
    if not essay:
        raise HTTPException(status_code=404, detail="作文不存在")

    if essay.content_file:
        dir_path = os.path.dirname(os.path.join(get_upload_dir(), essay.content_file))
        safe_name = os.path.basename(filename)
        file_path = os.path.join(dir_path, safe_name)
        if os.path.exists(file_path):
            import mimetypes
            media_type = mimetypes.guess_type(file_path)[0] or "application/octet-stream"
            return FileResponse(file_path, media_type=media_type,
                                headers={"Cache-Control": "no-cache"})

    db_img = db.query(EssayImage).filter(
        EssayImage.essay_id == essay_id,
        EssayImage.filename == filename,
    ).first()
    if db_img:
        import mimetypes
        media_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"
        from fastapi.responses import Response
        return Response(content=db_img.image_data, media_type=media_type,
                        headers={"Cache-Control": "no-cache"})

    raise HTTPException(status_code=404, detail="文件不存在")


@router.get("/{essay_id}/download")
def download_essay_file(
    essay_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """下载原文：有文字内容时从 DB 生成 docx，纯图片时打包 zip"""
    from starlette.background import BackgroundTask

    def _respond_with_cleanup(path, filename, media_type):
        def _cleanup():
            try:
                if os.path.exists(path):
                    os.remove(path)
            except OSError:
                pass
        return FileResponse(path, filename=filename, media_type=media_type,
                            background=BackgroundTask(_cleanup))
    if "guest" in current_user.role:
        raise HTTPException(status_code=403, detail="游客无下载权限")
    essay = db.query(Essay).filter(Essay.id == essay_id).first()
    if not essay:
        raise HTTPException(status_code=404, detail="作文不存在")

    dl_name = _build_download_filename(essay)

    # 收集图片文件（如果存在）
    img_files = []
    dir_path = ""
    if essay.content_file:
        dir_path = os.path.dirname(os.path.join(get_upload_dir(), essay.content_file))
        if os.path.exists(dir_path):
            all_files = os.listdir(dir_path)
            img_files = [f for f in all_files if f.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.webp')) and not f.startswith('改_')]

    has_text = bool(essay.content_text and essay.content_text.strip())

    # 既有文字又有图片 → docx + 图片打包 zip
    if has_text and img_files:
        tmp_docx = _generate_docx(essay, show_corrected=False)
        zip_buffer = tempfile.NamedTemporaryFile(delete=False, suffix=".zip")
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.write(tmp_docx, f"{dl_name}.docx")
            for img in sorted(img_files):
                zf.write(os.path.join(dir_path, img), img)
        zip_buffer.close()
        os.unlink(tmp_docx)
        return _respond_with_cleanup(zip_buffer.name, f"{dl_name}.zip", "application/zip")

    # 只有文字 → 从 DB 生成 docx
    if has_text:
        tmp_path = _generate_docx(essay, show_corrected=False)
        return _respond_with_cleanup(
            tmp_path,
            f"{dl_name}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    # 纯图片 → 打包 zip
    if essay.content_file and img_files:
        zip_buffer = tempfile.NamedTemporaryFile(delete=False, suffix=".zip")
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            for img in sorted(img_files):
                zf.write(os.path.join(dir_path, img), img)
        zip_buffer.close()
        return _respond_with_cleanup(zip_buffer.name, f"{dl_name}.zip", "application/zip")

    # 兜底：返回原始文件
    if essay.content_file:
        file_path = os.path.join(get_upload_dir(), essay.content_file)
        if os.path.exists(file_path):
            import mimetypes
            ext = os.path.splitext(file_path)[1]
            media_type = mimetypes.guess_type(file_path)[0] or "application/octet-stream"
            return FileResponse(file_path, filename=f"{dl_name}{ext}", media_type=media_type)

    raise HTTPException(status_code=404, detail="文件不存在")


@router.get("/{essay_id}/download-correction")
def download_correction(
    essay_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """下载修改结果"""
    essay = db.query(Essay).filter(Essay.id == essay_id).first()
    if not essay or not essay.content_file:
        raise HTTPException(status_code=404, detail="文件不存在")

    original_path = os.path.join(get_upload_dir(), essay.content_file)
    original_dir = os.path.dirname(original_path)
    original_name = os.path.basename(original_path)

    corr_name = generate_correction_filename(original_name)
    corr_path = os.path.join(original_dir, corr_name)

    if not os.path.exists(corr_path):
        raise HTTPException(status_code=404, detail="修改结果不存在")

    dl_name = _build_download_filename(essay)
    ext = os.path.splitext(corr_path)[1]
    return FileResponse(corr_path, filename=f"{dl_name}{ext}")


@router.get("/{essay_id}/export-docx")
def export_docx(
    essay_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """导出修改前后 docx：从 DB 读取 content_text + corrected_text"""
    if "guest" in current_user.role:
        raise HTTPException(status_code=403, detail="游客无导出权限")
    essay = db.query(Essay).filter(Essay.id == essay_id).first()
    if not essay:
        raise HTTPException(status_code=404, detail="作文不存在")

    tmp_path = _generate_docx(essay, show_corrected=True)
    dl_name = _build_download_filename(essay)

    from starlette.background import BackgroundTask

    def _cleanup_docx():
        try:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
        except OSError:
            pass

    return FileResponse(
        tmp_path,
        filename=f"改_{dl_name}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        background=BackgroundTask(_cleanup_docx),
    )


@router.post("/{essay_id}/ocr")
def ocr_essay(
    essay_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """对作文图片进行 OCR 识别，提取文字保存到 content_text"""
    if "collector" not in current_user.role and "reviewer" not in current_user.role and "admin" not in current_user.role:
        raise HTTPException(status_code=403, detail="无权限")
    essay = db.query(Essay).filter(Essay.id == essay_id).first()
    if not essay:
        raise HTTPException(status_code=404, detail="作文不存在")
    if essay.file_type != "image":
        raise HTTPException(status_code=400, detail="仅支持图片类型的作文进行 OCR")
    if not essay.content_file:
        raise HTTPException(status_code=400, detail="作文无文件")

    cfg_row = db.query(SystemConfig).filter(SystemConfig.config_key == "ocr").first()
    ocr_cfg = load_config_row_value(cfg_row.config_value) if cfg_row else {}
    if not ocr_cfg.get("enabled", False):
        raise HTTPException(status_code=400, detail="OCR 功能未启用，请先在系统设置中配置")

    xfyun_cfg = ocr_cfg.get("xfyun", {})
    if not xfyun_cfg.get("url") or not xfyun_cfg.get("appid") or not xfyun_cfg.get("api_key"):
        raise HTTPException(status_code=400, detail="讯飞 OCR 配置不完整")

    essay_dir = os.path.dirname(os.path.join(get_upload_dir(), essay.content_file))
    meta = {}
    try:
        text = ocr_essay_images_with_fallback(db, essay.id, essay_dir, xfyun_cfg, meta=meta)
        essay.content_text = text
        op_text = "OCR 识别完成"
        if meta.get("image_corrected"):
            op_text += f"（图片矫正 {meta['image_corrected']} 张，最大旋转 {meta['max_rotation']:.1f}°）"
        _log_operation(db, essay.id, current_user.id, "OCR", op_text)
        db.commit()
        db.refresh(essay)
        return {
            "content_text": text,
            "word_count": _count_non_ws(text),
            "image_corrected": meta.get("image_corrected", 0),
            "max_rotation": meta.get("max_rotation"),
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"OCR 识别失败: {str(e)}")


@router.post("/{essay_id}/ai-correct")
def ai_correct_essay(
    essay_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """对作文内容进行 AI 错别字修正，保存到 corrected_text"""
    if "collector" not in current_user.role and "reviewer" not in current_user.role and "admin" not in current_user.role:
        raise HTTPException(status_code=403, detail="无权限")
    essay = db.query(Essay).filter(Essay.id == essay_id).first()
    if not essay:
        raise HTTPException(status_code=404, detail="作文不存在")
    if not essay.content_text or not essay.content_text.strip():
        raise HTTPException(status_code=400, detail="作文无文字内容，请先进行 OCR 或手动输入")

    cfg_row = db.query(SystemConfig).filter(SystemConfig.config_key == "llm_typo_fix").first()
    if not cfg_row:
        raise HTTPException(status_code=400, detail="AI 错别字修正配置不存在，请先在系统设置中保存一次")
    try:
        llm_cfg = load_config_row_value(cfg_row.config_value) if cfg_row else {}
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="AI 错别字修正配置损坏，请重新保存系统设置")
    if not llm_cfg.get("enabled", False):
        raise HTTPException(status_code=400, detail="AI 错别字修正未启用，请在系统设置的「修改前-AI错别字修正」中勾选启用并保存")

    try:
        essay_info = {
            "student_name": essay.student_name,
            "grade": essay.grade,
            "essay_number": essay.essay_number,
            "essay_title": essay.essay_title,
            "teaching_mode": essay.teaching_mode,
            "task_name": essay.task.name if essay.task else None,
        }
        result = ai_correct_text(essay.content_text, llm_cfg, essay_info=essay_info)
        corrected_text = result.get("修改后内容", essay.content_text)
        essay.content_text = corrected_text
        if not essay.essay_title or not essay.essay_title.strip():
            title = result.get("作文标题", "")
            if title and title != "未知":
                essay.essay_title = title.strip()
        _log_operation(db, essay.id, current_user.id, "编辑", "AI 错别字修正")
        db.commit()
        db.refresh(essay)
        return {
            "content_text": corrected_text,
            "essay_title": essay.essay_title,
            "metadata": {
                "title": result.get("作文标题", "未知"),
                "author": result.get("作者", "未知"),
                "word_count": result.get("原文字数", "未知"),
                "grade": result.get("年级", "未知"),
                "mode": result.get("线上或线下", "未知"),
            },
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI 错别字修正失败: {str(e)}")


@router.post("/{essay_id}/ai-rewrite")
def ai_rewrite_essay(
    essay_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """对作文原文进行 AI 改写，结果保存到 corrected_text"""
    if "reviewer" not in current_user.role and "admin" not in current_user.role:
        raise HTTPException(status_code=403, detail="无权限")
    essay = db.query(Essay).filter(Essay.id == essay_id).first()
    if not essay:
        raise HTTPException(status_code=404, detail="作文不存在")
    if not essay.content_text or not essay.content_text.strip():
        raise HTTPException(status_code=400, detail="作文无文字内容")

    cfg_row = db.query(SystemConfig).filter(SystemConfig.config_key == "llm_editor").first()
    if not cfg_row:
        raise HTTPException(status_code=400, detail="AI 改作文配置不存在，请先在系统设置中保存一次")
    try:
        llm_cfg = load_config_row_value(cfg_row.config_value) if cfg_row else {}
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="AI 改作文配置损坏，请重新保存系统设置")
    if not llm_cfg.get("enabled", False):
        raise HTTPException(status_code=400, detail="AI 改作文未启用，请在系统设置的「修改后-AI改作文」中勾选启用并保存")

    count_min = llm_cfg.get("count_min")
    count_max = llm_cfg.get("count_max")
    if count_min is not None:
        try: count_min = int(count_min)
        except: count_min = None
    if count_max is not None:
        try: count_max = int(count_max)
        except: count_max = None

    try:
        rewritten = ai_rewrite_text(
            essay.content_text, llm_cfg,
            prompt_template=llm_cfg.get("prompt"),
            count_min=count_min, count_max=count_max,
        )
        essay.corrected_text = rewritten
        if essay.status in ("pending", "rework") and essay.content_text and essay.content_text.strip():
            essay.status = "confirming"
        essay.corrected_at = datetime.now()
        essay.reviewer_id = current_user.id
        _log_operation(db, essay.id, current_user.id, "批改", "AI 改写")
        db.commit()
        db.refresh(essay)
        return {"corrected_text": rewritten, "char_count": count_cjk_chars(rewritten)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI 改写失败: {str(e)}")


@router.post("/{essay_id}/confirm")
def confirm_essay(
    essay_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """确认修改：将作文从 待确认 改为 已修改"""
    if "reviewer" not in current_user.role and "admin" not in current_user.role:
        raise HTTPException(status_code=403, detail="无权限")
    essay = db.query(Essay).filter(Essay.id == essay_id).first()
    if not essay:
        raise HTTPException(status_code=404, detail="作文不存在")
    if essay.status != "confirming":
        raise HTTPException(status_code=400, detail="当前状态不是待确认，无法确认")
    essay.status = "corrected"
    essay.corrected_at = datetime.now()
    essay.reviewer_id = current_user.id
    _log_operation(db, essay.id, current_user.id, "批改", "确认修改")
    db.commit()
    db.refresh(essay)
    return _essay_to_out(essay, db)


@router.post("/{essay_id}/rework")
def rework_essay(
    essay_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """标记重改：将作文从 待确认 改为 待重改（修改后文章不达标，需重新改正）"""
    if "reviewer" not in current_user.role and "admin" not in current_user.role:
        raise HTTPException(status_code=403, detail="无权限")
    essay = db.query(Essay).filter(Essay.id == essay_id).first()
    if not essay:
        raise HTTPException(status_code=404, detail="作文不存在")
    if essay.status != "confirming":
        raise HTTPException(status_code=400, detail="当前状态不是待确认，无法标记为重改")
    essay.status = "rework"
    essay.reviewer_id = current_user.id
    _log_operation(db, essay.id, current_user.id, "批改", "标记为重改")
    db.commit()
    db.refresh(essay)
    return _essay_to_out(essay, db)


@router.post("/batch-update")
def batch_update_essays(
    data: dict,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """批量修改选中作文的收集者或任务"""
    if "admin" not in current_user.role:
        raise HTTPException(status_code=403, detail="仅管理员可批量修改")
    essay_ids = data.get("ids", [])
    if not essay_ids:
        raise HTTPException(status_code=400, detail="未选中任何作文")
    essays = db.query(Essay).filter(Essay.id.in_(essay_ids)).all()
    updated = 0
    skipped = 0
    if "collected_by" in data and data["collected_by"]:
        for e in essays:
            e.collected_by = data["collected_by"]
            updated += 1
    if "task_id" in data:
        from sqlalchemy.exc import IntegrityError
        new_task_id = data["task_id"] if data["task_id"] else None
        new_task = None
        if new_task_id:
            new_task = db.query(EssayTask).filter(EssayTask.id == new_task_id, EssayTask.deleted_at == None).first()
        for e in essays:
            savepoint = db.begin_nested()
            e.task_id = new_task_id
            # 更换任务时同步年级/第几次/课程为新任务的值（任务无值则清空）
            e.grade = (new_task.grade or "") if new_task else ""
            e.essay_number = (new_task.essay_number or 0) if new_task else 0
            e.course_id = (new_task.course_id or None) if new_task else None
            try:
                db.flush()
                savepoint.commit()
                updated += 1
            except IntegrityError:
                savepoint.rollback()
                skipped += 1
    db.commit()
    msg = f"已更新 {updated} 条记录，{skipped} 条略过（重复冲突）" if skipped else f"已更新 {updated} 条记录"
    return {"message": msg, "count": updated, "skipped": skipped}


@router.post("/batch-export-docx")
def batch_export_docx(
    essay_ids: list[int],
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """批量导出选中作文的docx（修改前后），打包为zip下载"""
    if "guest" in current_user.role:
        raise HTTPException(status_code=403, detail="游客无导出权限")
    from pydantic import BaseModel

    essays = db.query(Essay).filter(Essay.id.in_(essay_ids)).all()
    if not essays:
        raise HTTPException(status_code=404, detail="未找到选中的作文")

    # 创建临时zip文件
    tmp_zip = tempfile.NamedTemporaryFile(delete=False, suffix=".zip")
    tmp_zip_path = tmp_zip.name
    tmp_zip.close()

    try:
        with zipfile.ZipFile(tmp_zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            for essay in essays:
                tmp_docx = _generate_docx(essay, show_corrected=True)
                dl_name = _build_download_filename(essay)
                # 将docx文件添加到zip中
                zf.write(tmp_docx, f"改_{dl_name}.docx")
                # 删除临时docx文件
                os.unlink(tmp_docx)

        # 构建下载文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        zip_filename = f"作文导出_{timestamp}.zip"

        from starlette.background import BackgroundTask

        def _cleanup_zip():
            try:
                if os.path.exists(tmp_zip_path):
                    os.unlink(tmp_zip_path)
            except OSError:
                pass

        return FileResponse(
            tmp_zip_path,
            filename=zip_filename,
            media_type="application/zip",
            background=BackgroundTask(_cleanup_zip),
        )
    except Exception as e:
        # 清理临时文件
        if os.path.exists(tmp_zip_path):
            os.unlink(tmp_zip_path)
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


@router.post("/batch-export-docx-merged")
def batch_export_docx_merged(
    essay_ids: list[int],
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """一键合并修改前后 docx：把选中作文的修改前后内容合并为一个 docx。
    文件名使用任务名称，存在多个任务时按多数任务名称命名。"""
    if "guest" in current_user.role:
        raise HTTPException(status_code=403, detail="游客无导出权限")
    from collections import Counter
    from docx import Document

    essays = db.query(Essay).filter(Essay.id.in_(essay_ids)).all()
    if not essays:
        raise HTTPException(status_code=404, detail="未找到选中的作文")
    if len(essays) > 200:
        raise HTTPException(status_code=400, detail="一次最多合并 200 篇作文，请分批导出")

    # 统计各任务下的作文数，取多数任务名称
    task_ids = {e.task_id for e in essays if e.task_id}
    task_names = {}
    if task_ids:
        for t in db.query(EssayTask).filter(EssayTask.id.in_(task_ids)).all():
            task_names[t.id] = t.name or "未命名任务"

    counter = Counter()
    for e in essays:
        name = task_names.get(e.task_id) if e.task_id else ""
        counter[name or "未关联任务"] += 1
    majority_name = counter.most_common(1)[0][0] if counter else "作文合并"
    safe_name = majority_name.replace("/", "_").replace("\\", "_").strip() or "作文合并"

    doc = Document()
    for idx, essay in enumerate(essays):
        if idx > 0:
            doc.add_page_break()
        _append_essay_to_doc(doc, essay, show_corrected=True)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_path = tmp.name
    tmp.close()
    doc.save(tmp_path)

    from starlette.background import BackgroundTask

    def _cleanup_merged_docx():
        try:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
        except OSError:
            pass

    return FileResponse(
        tmp_path,
        filename=f"{safe_name}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        background=BackgroundTask(_cleanup_merged_docx),
    )


@router.post("/batch-export-docx-corrected-merged")
def batch_export_docx_corrected_merged(
    essay_ids: list[int],
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """一键合并仅修改后 docx：只输出选中作文的修改后内容为一个 docx。
    文件名使用任务名称，存在多个任务时按多数任务名称命名。"""
    if "guest" in current_user.role:
        raise HTTPException(status_code=403, detail="游客无导出权限")
    from collections import Counter
    from docx import Document

    essays = db.query(Essay).filter(Essay.id.in_(essay_ids)).all()
    if not essays:
        raise HTTPException(status_code=404, detail="未找到选中的作文")
    if len(essays) > 200:
        raise HTTPException(status_code=400, detail="一次最多合并 200 篇作文，请分批导出")

    task_ids = {e.task_id for e in essays if e.task_id}
    task_names = {}
    if task_ids:
        for t in db.query(EssayTask).filter(EssayTask.id.in_(task_ids)).all():
            task_names[t.id] = t.name or "未命名任务"

    counter = Counter()
    for e in essays:
        name = task_names.get(e.task_id) if e.task_id else ""
        counter[name or "未关联任务"] += 1
    majority_name = counter.most_common(1)[0][0] if counter else "作文合并"
    safe_name = majority_name.replace("/", "_").replace("\\", "_").strip() or "作文合并"

    doc = Document()
    for idx, essay in enumerate(essays):
        if idx > 0:
            doc.add_page_break()
        _append_essay_to_doc(doc, essay, show_corrected=True, add_heading=True, corrected_only=True)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_path = tmp.name
    tmp.close()
    doc.save(tmp_path)

    from starlette.background import BackgroundTask

    def _cleanup_merged_docx():
        try:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
        except OSError:
            pass

    return FileResponse(
        tmp_path,
        filename=f"{safe_name}_修改后.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        background=BackgroundTask(_cleanup_merged_docx),
    )


@router.post("/batch-export-docx-original-merged")
def batch_export_docx_original_merged(
    essay_ids: list[int],
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """一键合并仅修改前 docx：只输出选中作文的修改前（原文）内容为一个 docx。
    文件名使用任务名称，存在多个任务时按多数任务名称命名。"""
    if "guest" in current_user.role:
        raise HTTPException(status_code=403, detail="游客无导出权限")
    from collections import Counter
    from docx import Document

    essays = db.query(Essay).filter(Essay.id.in_(essay_ids)).all()
    if not essays:
        raise HTTPException(status_code=404, detail="未找到选中的作文")
    if len(essays) > 200:
        raise HTTPException(status_code=400, detail="一次最多合并 200 篇作文，请分批导出")

    task_ids = {e.task_id for e in essays if e.task_id}
    task_names = {}
    if task_ids:
        for t in db.query(EssayTask).filter(EssayTask.id.in_(task_ids)).all():
            task_names[t.id] = t.name or "未命名任务"

    counter = Counter()
    for e in essays:
        name = task_names.get(e.task_id) if e.task_id else ""
        counter[name or "未关联任务"] += 1
    majority_name = counter.most_common(1)[0][0] if counter else "作文合并"
    safe_name = majority_name.replace("/", "_").replace("\\", "_").strip() or "作文合并"

    doc = Document()
    for idx, essay in enumerate(essays):
        if idx > 0:
            doc.add_page_break()
        _append_essay_to_doc(doc, essay, add_heading=True, original_only=True)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_path = tmp.name
    tmp.close()
    doc.save(tmp_path)

    from starlette.background import BackgroundTask

    def _cleanup_merged_docx():
        try:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
        except OSError:
            pass

    return FileResponse(
        tmp_path,
        filename=f"{safe_name}_修改前.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        background=BackgroundTask(_cleanup_merged_docx),
    )


@router.post("/batch-ocr")
def batch_ocr_essays(
    data: dict,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """批量 OCR 识别选中作文的图片"""
    if "reviewer" not in current_user.role and "admin" not in current_user.role:
        raise HTTPException(status_code=403, detail="无权限")
    essay_ids = data.get("ids", [])
    if not essay_ids:
        raise HTTPException(status_code=400, detail="未选中任何作文")
    essays = db.query(Essay).filter(Essay.id.in_(essay_ids), Essay.deleted_at == None).all()
    if not essays:
        raise HTTPException(status_code=404, detail="未找到选中的作文")

    cfg_row = db.query(SystemConfig).filter(SystemConfig.config_key == "ocr").first()
    ocr_cfg = load_config_row_value(cfg_row.config_value) if cfg_row else {}
    if not ocr_cfg.get("enabled", False):
        raise HTTPException(status_code=400, detail="OCR 功能未启用，请先在系统设置中配置")
    xfyun_cfg = ocr_cfg.get("xfyun", {})
    if not xfyun_cfg.get("url") or not xfyun_cfg.get("appid") or not xfyun_cfg.get("api_key"):
        raise HTTPException(status_code=400, detail="讯飞 OCR 配置不完整")

    success = 0
    errors = []
    for e in essays:
        if e.file_type != "image" or not e.content_file:
            errors.append({"id": e.id, "student": e.student_name, "reason": "非图片类型或无文件"})
            continue
        try:
            essay_dir = os.path.dirname(os.path.join(get_upload_dir(), e.content_file))
            meta = {}
            text = ocr_essay_images_with_fallback(db, e.id, essay_dir, xfyun_cfg, meta=meta)
            e.content_text = text
            op_text = "批量 OCR 识别完成"
            if meta.get("image_corrected"):
                op_text += f"（图片矫正 {meta['image_corrected']} 张，最大旋转 {meta['max_rotation']:.1f}°）"
            _log_operation(db, e.id, current_user.id, "OCR", op_text)
            success += 1
        except Exception as ex:
            errors.append({"id": e.id, "student": e.student_name, "reason": str(ex)})
    db.commit()
    return {"success": success, "errors": errors, "total": len(essays)}


@router.post("/batch-ai-correct")
def batch_ai_correct_essays(
    data: dict,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """批量 AI 错别字修正选中作文的内容"""
    if "reviewer" not in current_user.role and "admin" not in current_user.role:
        raise HTTPException(status_code=403, detail="无权限")
    essay_ids = data.get("ids", [])
    if not essay_ids:
        raise HTTPException(status_code=400, detail="未选中任何作文")
    essays = db.query(Essay).filter(Essay.id.in_(essay_ids), Essay.deleted_at == None).all()
    if not essays:
        raise HTTPException(status_code=404, detail="未找到选中的作文")

    cfg_row = db.query(SystemConfig).filter(SystemConfig.config_key == "llm_typo_fix").first()
    if not cfg_row:
        raise HTTPException(status_code=400, detail="AI 错别字修正配置不存在，请先在系统设置中保存一次")
    try:
        llm_cfg = load_config_row_value(cfg_row.config_value) if cfg_row else {}
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="AI 错别字修正配置损坏，请重新保存系统设置")
    if not llm_cfg.get("enabled", False):
        raise HTTPException(status_code=400, detail="AI 错别字修正未启用，请在系统设置的「修改前-AI错别字修正」中勾选启用并保存")

    success = 0
    errors = []
    for e in essays:
        if not e.content_text or not e.content_text.strip():
            errors.append({"id": e.id, "student": e.student_name, "reason": "无文字内容"})
            continue
        try:
            result = ai_correct_text(e.content_text, llm_cfg)
            corrected_text = result.get("修改后内容", e.content_text)
            e.content_text = corrected_text
            _log_operation(db, e.id, current_user.id, "编辑", "批量 AI 错别字修正")
            success += 1
        except Exception as ex:
            errors.append({"id": e.id, "student": e.student_name, "reason": str(ex)})
    db.commit()
    return {"success": success, "errors": errors, "total": len(essays)}


@router.post("/batch-ai-rewrite")
def batch_ai_rewrite_essays(
    data: dict,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """批量 AI 改写（一键修改）选中作文，结果保存到 corrected_text"""
    if "reviewer" not in current_user.role and "admin" not in current_user.role:
        raise HTTPException(status_code=403, detail="无权限")
    essay_ids = data.get("ids", [])
    if not essay_ids:
        raise HTTPException(status_code=400, detail="未选中任何作文")
    essays = db.query(Essay).filter(Essay.id.in_(essay_ids), Essay.deleted_at == None).all()
    if not essays:
        raise HTTPException(status_code=404, detail="未找到选中的作文")

    cfg_row = db.query(SystemConfig).filter(SystemConfig.config_key == "llm_editor").first()
    if not cfg_row:
        raise HTTPException(status_code=400, detail="AI 改作文配置不存在，请先在系统设置中保存一次")
    try:
        llm_cfg = load_config_row_value(cfg_row.config_value) if cfg_row else {}
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="AI 改作文配置损坏，请重新保存系统设置")
    if not llm_cfg.get("enabled", False):
        raise HTTPException(status_code=400, detail="AI 改作文未启用，请在系统设置的「修改后-AI改作文」中勾选启用并保存")

    success = 0
    errors = []
    for e in essays:
        if not e.content_text or not e.content_text.strip():
            errors.append({"id": e.id, "student": e.student_name, "reason": "无文字内容"})
            continue
        try:
            rewritten = ai_rewrite_text(e.content_text, llm_cfg, prompt_template=llm_cfg.get("prompt"))
            e.corrected_text = rewritten
            if e.status in ("pending", "rework") and e.content_text and e.content_text.strip():
                e.status = "confirming"
            e.corrected_at = datetime.now()
            e.reviewer_id = current_user.id
            _log_operation(db, e.id, current_user.id, "批改", "批量 AI 改写")
            success += 1
        except Exception as ex:
            errors.append({"id": e.id, "student": e.student_name, "reason": str(ex)})
    db.commit()
    return {"success": success, "errors": errors, "total": len(essays)}


@router.post("/batch-confirm")
def batch_confirm_essays(
    data: dict,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """批量确认修改：将选中作文从 待确认 改为 已修改"""
    if "reviewer" not in current_user.role and "admin" not in current_user.role:
        raise HTTPException(status_code=403, detail="无权限")
    essay_ids = data.get("ids", [])
    if not essay_ids:
        raise HTTPException(status_code=400, detail="未选中任何作文")
    essays = db.query(Essay).filter(Essay.id.in_(essay_ids), Essay.deleted_at == None).all()
    if not essays:
        raise HTTPException(status_code=404, detail="未找到选中的作文")
    count = 0
    for e in essays:
        if e.status == "confirming":
            e.status = "corrected"
            e.corrected_at = datetime.now()
            e.reviewer_id = current_user.id
            _log_operation(db, e.id, current_user.id, "批改", "确认修改")
            count += 1
    db.commit()
    return {"success": count, "total": len(essays)}


# ===== 异步批量任务（后台运行，页面离开后继续执行） =====

def _get_ocr_config(db):
    cfg_row = db.query(SystemConfig).filter(SystemConfig.config_key == "ocr").first()
    if cfg_row:
        try:
            return load_config_row_value(cfg_row.config_value)
        except json.JSONDecodeError:
            return {}
    return {}

def _get_llm_config(db, key):
    cfg_row = db.query(SystemConfig).filter(SystemConfig.config_key == key).first()
    if cfg_row:
        try:
            return load_config_row_value(cfg_row.config_value)
        except json.JSONDecodeError:
            return {}
    return {}


@router.post("/batch-task/ocr/start")
def start_batch_ocr(
    data: dict,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """启动异步批量 OCR 任务"""
    if "admin" not in current_user.role and "reviewer" not in current_user.role:
        raise HTTPException(status_code=403, detail="仅管理员或批改者可执行")
    essay_ids = data.get("ids", [])
    if not essay_ids:
        raise HTTPException(status_code=400, detail="未选中任何作文")
    ocr_cfg = _get_ocr_config(db)
    if not ocr_cfg.get("enabled", False):
        raise HTTPException(status_code=400, detail="OCR 功能未启用")

    task_id = str(uuid.uuid4())[:8]
    essays = db.query(Essay).filter(Essay.id.in_(essay_ids), Essay.deleted_at == None).all()
    from ..services.task_manager import create_task, run_batch_ocr, update_task

    task = create_task(task_id, "ocr", len(essays))
    thread = threading.Thread(target=run_batch_ocr, args=(
        task_id, essay_ids, current_user.id, ocr_cfg,
        get_db, Essay, _log_operation,
    ), daemon=True)
    thread.start()
    return {"task_id": task_id, "total": len(essays)}


@router.post("/batch-task/ai-correct/start")
def start_batch_ai_correct(
    data: dict,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """启动异步批量 AI 错别字修正"""
    if "admin" not in current_user.role and "reviewer" not in current_user.role:
        raise HTTPException(status_code=403, detail="仅管理员或批改者可执行")
    essay_ids = data.get("ids", [])
    if not essay_ids:
        raise HTTPException(status_code=400, detail="未选中任何作文")
    llm_cfg = _get_llm_config(db, "llm_typo_fix")
    if not llm_cfg.get("enabled", False):
        raise HTTPException(status_code=400, detail="AI 错别字修正未启用")

    task_id = str(uuid.uuid4())[:8]
    essays = db.query(Essay).filter(Essay.id.in_(essay_ids), Essay.deleted_at == None).all()
    from ..services.task_manager import create_task, run_batch_ai_correct

    create_task(task_id, "ai_correct", len(essays))
    thread = threading.Thread(target=run_batch_ai_correct, args=(
        task_id, essay_ids, current_user.id, llm_cfg,
        get_db, Essay, _log_operation,
    ), daemon=True)
    thread.start()
    return {"task_id": task_id, "total": len(essays)}


@router.post("/batch-task/ai-rewrite/start")
def start_batch_ai_rewrite(
    data: dict,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """启动异步批量 AI 一键修改"""
    if "admin" not in current_user.role and "reviewer" not in current_user.role:
        raise HTTPException(status_code=403, detail="仅管理员或批改者可执行")
    essay_ids = data.get("ids", [])
    if not essay_ids:
        raise HTTPException(status_code=400, detail="未选中任何作文")
    llm_cfg = _get_llm_config(db, "llm_editor")
    if not llm_cfg.get("enabled", False):
        raise HTTPException(status_code=400, detail="AI 改作文未启用")

    task_id = str(uuid.uuid4())[:8]
    essays = db.query(Essay).filter(Essay.id.in_(essay_ids), Essay.deleted_at == None).all()
    from ..services.task_manager import create_task, run_batch_ai_rewrite

    create_task(task_id, "ai_rewrite", len(essays))
    thread = threading.Thread(target=run_batch_ai_rewrite, args=(
        task_id, essay_ids, current_user.id, llm_cfg,
        get_db, Essay, _log_operation,
    ), daemon=True)
    thread.start()
    return {"task_id": task_id, "total": len(essays)}


@router.post("/batch-task/pipeline/start")
def start_batch_pipeline(
    data: dict,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """启动后台流水线：OCR → AI错别字修正 → AI一键修改（仅处理未修改的作文）"""
    if "admin" not in current_user.role and "reviewer" not in current_user.role:
        raise HTTPException(status_code=403, detail="仅管理员或批改者可执行")
    essay_ids = data.get("ids", [])
    if not essay_ids:
        raise HTTPException(status_code=400, detail="未选中任何作文")

    ocr_cfg = _get_ocr_config(db)
    if not ocr_cfg.get("enabled", False):
        raise HTTPException(status_code=400, detail="OCR 功能未启用")
    typo_cfg = _get_llm_config(db, "llm_typo_fix")
    if not typo_cfg.get("enabled", False):
        raise HTTPException(status_code=400, detail="AI 错别字修正未启用")
    editor_cfg = _get_llm_config(db, "llm_editor")
    if not editor_cfg.get("enabled", False):
        raise HTTPException(status_code=400, detail="AI 改作文未启用")

    base = str(uuid.uuid4())[:8]
    essays = db.query(Essay).filter(Essay.id.in_(essay_ids), Essay.deleted_at == None, Essay.status == "pending").all()
    total = len(essays)
    if not total:
        raise HTTPException(status_code=400, detail="选中的条目中没有状态为「未修改」的作文")
    pending_ids = [e.id for e in essays]

    from ..services.task_manager import create_task, run_batch_pipeline

    ocr_task_id = f"{base}-ocr"
    correct_task_id = f"{base}-correct"
    rewrite_task_id = f"{base}-rewrite"
    create_task(ocr_task_id, "ocr", total)
    create_task(correct_task_id, "ai_correct", total)
    create_task(rewrite_task_id, "ai_rewrite", total)
    thread = threading.Thread(target=run_batch_pipeline, args=(
        ocr_task_id, correct_task_id, rewrite_task_id, pending_ids,
        current_user.id, ocr_cfg, typo_cfg, editor_cfg,
        get_db, Essay, _log_operation,
    ), daemon=True)
    thread.start()
    return {
        "total": total,
        "tasks": [
            {"id": ocr_task_id, "type": "ocr", "total": total},
            {"id": correct_task_id, "type": "ai_correct", "total": total},
            {"id": rewrite_task_id, "type": "ai_rewrite", "total": total},
        ],
    }


@router.get("/batch-task/{task_id}")
def get_batch_task_status(
    task_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """查询批量任务进度"""
    from ..services.task_manager import get_task
    task = get_task(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在或已过期")
    return task


@router.post("/operations/{log_id}/undo")
def undo_operation(
    log_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """撤回指定的操作记录（仅管理员）"""
    if "admin" not in current_user.role:
        raise HTTPException(status_code=403, detail="无权限")

    log = db.query(OperationLog).filter(OperationLog.id == log_id).first()
    if not log:
        raise HTTPException(status_code=404, detail="操作记录不存在")

    action_str = log.action.value if hasattr(log.action, 'value') else str(log.action)
    undone_count = 0
    essay_ids_list = []

    # 解析批量作文ID
    if log.essay_ids:
        try:
            essay_ids_list = json.loads(log.essay_ids)
        except Exception:
            essay_ids_list = []
    elif log.essay_id:
        essay_ids_list = [log.essay_id]

    for eid in essay_ids_list:
        essay = db.query(Essay).filter(Essay.id == eid).first()
        if not essay:
            continue

        if action_str in ("删除", "DELETE"):
            essay.deleted_at = None
            _log_operation(db, eid, current_user.id, "恢复",
                           f"撤回删除操作", batch_id=log.batch_id)
            undone_count += 1

        elif action_str in ("恢复", "RECOVER"):
            essay.deleted_at = datetime.now()
            _log_operation(db, eid, current_user.id, "删除",
                           f"撤回恢复操作", batch_id=log.batch_id)
            undone_count += 1

        elif action_str in ("上传", "UPLOAD"):
            essay.deleted_at = datetime.now()
            _log_operation(db, eid, current_user.id, "删除",
                           f"撤回上传操作", batch_id=log.batch_id)
            undone_count += 1

        elif action_str in ("修改", "UPDATE", "批改", "CORRECT"):
            data = {}
            if log.old_value:
                try:
                    data = json.loads(log.old_value)
                except Exception:
                    data = {"corrected_text": "", "status": "pending"}
            essay.corrected_text = data.get("corrected_text", "")
            essay.corrected_at = None
            essay.reviewer_id = None
            essay.status = "pending"
            _log_operation(db, eid, current_user.id, "编辑",
                           f"撤回修改操作", batch_id=log.batch_id)
            undone_count += 1

        elif action_str in ("编辑", "EDIT"):
            if log.old_value:
                try:
                    data = json.loads(log.old_value)
                    for field, val in data.items():
                        if hasattr(essay, field) and isinstance(val, dict) and "old" in val:
                            setattr(essay, field, val["old"])
                except Exception:
                    pass
            _log_operation(db, eid, current_user.id, "编辑",
                           f"撤回编辑操作", batch_id=log.batch_id)
            undone_count += 1

        elif action_str in ("OCR",):
            essay.content_text = ""
            _log_operation(db, eid, current_user.id, "编辑",
                           f"撤回OCR操作", batch_id=log.batch_id)
            undone_count += 1

    db.commit()
    return {"message": f"已撤回 {undone_count} 条", "undone_count": undone_count}


@router.post("/batch-upload")
async def batch_upload_essays(
    grade: str = Form(...),
    essay_number: int = Form(...),
    teaching_mode: str = Form("线下"),
    data_list: str = Form(...),
    file: UploadFile = File(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """批量上传：在单个请求中上传多个作文，记录为一条操作日志"""
    if "collector" not in current_user.role and "admin" not in current_user.role:
        raise HTTPException(status_code=403, detail="无权限")

    import uuid as _uuid
    batch_uuid = _uuid.uuid4().hex[:12]

    try:
        items = json.loads(data_list)
    except Exception:
        raise HTTPException(status_code=400, detail="数据格式错误，需要 JSON 数组")

    now = datetime.now()
    safe_grade = safe_component(grade, "")
    safe_mode = safe_component(teaching_mode, "")
    grade_name = f"{safe_grade}{safe_mode}" if safe_mode else (safe_grade or "未定年级")

    dir_path = os.path.join(
        get_upload_dir(), safe_component(str(now.year), "0000"),
        safe_component(f"{now.month}月", "1月"), safe_component(str(now.day), "1"),
        f"{grade_name}第{essay_number}次" if essay_number else grade_name,
    )
    os.makedirs(dir_path, exist_ok=True)

    if file and file.filename:
        safe_filename = os.path.basename(file.filename)
        file_path = os.path.join(dir_path, safe_filename)
        content = await file.read()
        with open(file_path, "wb") as fw:
            fw.write(content)

    created_ids = []
    skipped_students = []

    task_id_from_item = None
    if items:
        task_id_from_item = items[0].get("task_id")

    existing_names = set()
    if task_id_from_item:
        existing = db.query(Essay.student_name).filter(
            Essay.task_id == task_id_from_item,
            Essay.grade == grade,
            Essay.essay_number == essay_number,
            Essay.deleted_at == None,
        ).all()
        existing_names = {row[0] for row in existing}

    seen_in_batch = set()
    for item in items:
        student_name = item.get("student_name", "")
        if student_name in existing_names or student_name in seen_in_batch:
            skipped_students.append(student_name)
            continue
        savepoint = db.begin_nested()
        try:
            essay = Essay(
                grade=grade,
                essay_number=essay_number,
                essay_title=item.get("essay_title", ""),
                student_name=student_name,
                is_supplement=item.get("is_supplement", False),
                teaching_mode=teaching_mode,
                remark=item.get("remark", ""),
                content_text=item.get("content_text", ""),
                corrected_text=item.get("corrected_text", ""),
                file_type="docx",
                collected_by=current_user.id,
                task_id=task_id_from_item,
                status="pending",
            )
            db.add(essay)
            db.flush()
            savepoint.commit()
            created_ids.append(essay.id)
            seen_in_batch.add(student_name)
        except IntegrityError:
            savepoint.rollback()
            skipped_students.append(student_name)
            continue

    db.commit()

    detail_text = f"批量上传 {len(created_ids)} 篇 ({grade}第{essay_number}次)"
    if skipped_students:
        detail_text += f"，跳过 {len(skipped_students)} 个已存在学生：{'、'.join(skipped_students)}"

    # 插入批量操作日志
    try:
        batch_log = OperationLog(
            essay_id=None,
            user_id=current_user.id,
            action="上传",
            detail=detail_text,
            batch_id=batch_uuid,
            essay_ids=json.dumps(created_ids),
        )
        db.add(batch_log)
        db.commit()
    except Exception:
        pass

    return {"message": detail_text, "count": len(created_ids), "ids": created_ids, "batch_id": batch_uuid, "skipped": skipped_students}


# ===== /{essay_id} 通用路由必须放在所有具名路由之后 =====


@router.get("/operations")
def list_operations(
    page: int = 1,
    page_size: int = 50,
    keyword: str = None,
    action: str = None,
    user_id: int = None,
    student_name: str = None,
    date_from: str = None,
    date_to: str = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """获取操作历史列表，支持按关键词/操作类型/操作者/学生/日期筛选。"""
    q = db.query(OperationLog)
    if keyword:
        kw = f"%{keyword}%"
        q = q.filter(
            (OperationLog.detail.like(kw)) | (OperationLog.action.like(kw))
        )
    if action:
        q = q.filter(OperationLog.action == action)
    if user_id:
        q = q.filter(OperationLog.user_id == user_id)
    if student_name:
        q = q.join(Essay, OperationLog.essay_id == Essay.id, isouter=True).filter(
            Essay.student_name.like(f"%{student_name}%")
        )
    if date_from:
        q = q.filter(OperationLog.created_at >= date_from)
    if date_to:
        q = q.filter(OperationLog.created_at <= date_to + " 23:59:59")
    q = q.order_by(OperationLog.created_at.desc())
    total = q.count()
    logs = q.offset((page - 1) * page_size).limit(page_size).all()

    result = []
    for log in logs:
        user = db.query(User).filter(User.id == log.user_id).first()
        essay = db.query(Essay).filter(Essay.id == log.essay_id).first() if log.essay_id else None
        result.append(OperationLogOut(
            id=log.id,
            essay_id=log.essay_id,
            user_id=log.user_id,
            user_name=user.nickname or user.username if user else "未知",
            action=log.action.value if hasattr(log.action, 'value') else log.action,
            old_value=log.old_value or "",
            new_value=log.new_value or "",
            detail=log.detail or "",
            batch_id=log.batch_id,
            essay_ids=log.essay_ids,
            student_name=essay.student_name if essay else "",
            essay_title=essay.essay_title if essay else "",
            essay_number=essay.essay_number if essay else 0,
            created_at=log.created_at,
        ))

    return {"items": result, "total": total, "page": page, "page_size": page_size}


@router.get("/{essay_id}", response_model=EssayOut)
def get_essay(
    essay_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    essay = db.query(Essay).filter(Essay.id == essay_id).first()
    if not essay:
        raise HTTPException(status_code=404, detail="作文不存在")
    return _essay_to_out(essay, db)


@router.put("/{essay_id}", response_model=EssayOut)
def update_essay(
    essay_id: int,
    grade: str = "",
    essay_number: int = None,
    essay_title: str = "",
    student_name: str = "",
    teaching_mode: str = "",
    remark: str = "",
    collector_note: str = None,
    reviewer_note: str = None,
    collected_by: int = None,
    is_supplement: bool = None,
    task_id: int = None,
    content_text: str = None,
    corrected_text: str = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """更新作文信息"""
    essay = db.query(Essay).filter(Essay.id == essay_id).first()
    if not essay:
        raise HTTPException(status_code=404, detail="作文不存在")
    # 收集者可修改大部分字段，批改者可修改 reviewer_note，修改后内容仅批改者/管理员可改
    can_edit = "admin" in current_user.role or essay.collected_by == current_user.id
    can_edit_review_note = can_edit or essay.reviewer_id == current_user.id
    can_edit_corrected = "reviewer" in current_user.role or "admin" in current_user.role
    if not can_edit and not can_edit_review_note and not can_edit_corrected:
        raise HTTPException(status_code=403, detail="无权限编辑此作文")

    # 记录修改前的路径关键字段
    old_grade = essay.grade
    old_number = essay.essay_number
    old_student = essay.student_name
    old_mode = essay.teaching_mode
    old_title = essay.essay_title or ""
    old_supplement = essay.is_supplement or False

    if grade and can_edit:
        essay.grade = grade
    if essay_number is not None and can_edit:
        essay.essay_number = essay_number
    if essay_title and can_edit:
        essay.essay_title = essay_title
    if student_name and can_edit:
        essay.student_name = student_name
    if teaching_mode and can_edit:
        essay.teaching_mode = teaching_mode
    if remark is not None and can_edit:
        essay.remark = remark
    if collected_by is not None and "admin" in current_user.role:
        essay.collected_by = collected_by
    if is_supplement is not None and can_edit:
        essay.is_supplement = is_supplement
    if task_id is not None and can_edit:
        essay.task_id = task_id if task_id > 0 else None
    if content_text is not None and can_edit:
        essay.content_text = content_text
    if corrected_text is not None and can_edit_corrected:
        essay.corrected_text = corrected_text
    if collector_note is not None and can_edit:
        essay.collector_note = collector_note
    if reviewer_note is not None and can_edit_review_note:
        essay.reviewer_note = reviewer_note

    # 如果年级/次数/学生/方式变了，移动文件
    new_grade = essay.grade
    new_number = essay.essay_number
    new_student = essay.student_name
    new_mode = essay.teaching_mode

    if (essay.content_file and (
        new_grade != old_grade or
        new_number != old_number or
        new_student != old_student or
        new_mode != old_mode or
        (essay.essay_title or "") != old_title or
        (essay.is_supplement or False) != old_supplement
    )):
        old_dir = os.path.dirname(os.path.join(get_upload_dir(), essay.content_file))

        # 构建新目录路径
        from ..utils.file_utils import get_essay_dir, move_content_file
        now = datetime.now()
        task_name = ""
        task_created_at = None
        if essay.task_id:
            task = db.query(EssayTask).filter(EssayTask.id == essay.task_id).first()
            if task:
                task_name = task.name
                task_created_at = task.created_at

        new_dir = get_essay_dir(
            str(now.year), f"{now.month}月", str(now.day),
            new_grade or "未定年级", new_number, "", new_student, new_mode,
            task_name=task_name, task_created_at=task_created_at,
            essay_title=essay.essay_title, is_supplement=essay.is_supplement,
        )

        new_content_file = move_content_file(essay, old_dir, new_dir, filenames=_essay_owned_filenames(essay, db))
        if new_content_file:
            essay.content_file = new_content_file

    # 构建变更记录
    import json
    changes = {}
    if old_grade != essay.grade:
        changes["grade"] = {"old": old_grade, "new": essay.grade}
    if old_number != essay.essay_number:
        changes["essay_number"] = {"old": old_number, "new": essay.essay_number}
    if old_student != essay.student_name:
        changes["student_name"] = {"old": old_student, "new": essay.student_name}
    if old_mode != essay.teaching_mode:
        changes["teaching_mode"] = {"old": old_mode, "new": essay.teaching_mode}
    if essay_title and essay_title != essay.essay_title:
        changes["essay_title"] = {"old": essay.essay_title, "new": essay_title}
    if remark is not None and remark != essay.remark:
        changes["remark"] = {"old": essay.remark, "new": remark}
    if is_supplement is not None and is_supplement != essay.is_supplement:
        changes["is_supplement"] = {"old": essay.is_supplement, "new": is_supplement}

    old_value = json.dumps(changes, ensure_ascii=False) if changes else ""
    _log_operation(db, essay.id, current_user.id, "编辑", essay.student_name, old_value=old_value, new_value=old_value)
    try:
        db.commit()
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"保存失败: {str(e)}")
    db.refresh(essay)
    return _essay_to_out(essay, db)


def _essay_to_out(essay: Essay, db: Session) -> EssayOut:
    collector = db.query(User).filter(User.id == essay.collected_by).first()
    reviewer = db.query(User).filter(User.id == essay.reviewer_id).first() if essay.reviewer_id else None
    task = db.query(EssayTask).filter(EssayTask.id == essay.task_id).first() if essay.task_id else None
    course = db.query(Course).filter(Course.id == essay.course_id).first() if essay.course_id else None

    corr_exists = False
    file_path = ""
    if essay.content_file:
        file_path = os.path.join(get_upload_dir(), essay.content_file)
        original_dir = os.path.dirname(file_path)
        original_name = os.path.basename(file_path)
        corr_exists = has_correction(original_dir, original_name)

    # 自动同步状态
    if corr_exists and essay.status in ("pending", "rework") and essay.content_text and essay.content_text.strip():
        essay.status = "confirming"
        db.commit()

    return EssayOut(
        id=essay.id,
        task_id=essay.task_id,
        task_name=task.name if task else "",
        course_id=essay.course_id,
        course_name=course.name if course else (task.course_name if task else ""),
        grade=essay.grade or "",
        essay_number=essay.essay_number or 0,
        essay_title=essay.essay_title or "",
        student_name=essay.student_name,
        is_supplement=essay.is_supplement or False,
        teaching_mode=essay.teaching_mode or "线下",
        remark=essay.remark or "",
        collector_note=essay.collector_note or "",
        reviewer_note=essay.reviewer_note or "",
        content_text=essay.content_text or "",
        corrected_text=essay.corrected_text or "",
        content_file=essay.content_file or "",
        file_type=essay.file_type or "text",
        collected_by=essay.collected_by,
        collector_name=collector.nickname or collector.username if collector else "未知",
        status=essay.status or "pending",
        reviewer_id=essay.reviewer_id,
        reviewer_name=reviewer.nickname or reviewer.username if reviewer else "",
        corrected_at=essay.corrected_at,
        created_at=essay.created_at,
        file_path=file_path,
        has_correction=corr_exists,
        file_saved=essay.file_saved if essay.file_saved is not None else True,
        word_count=_count_non_ws(essay.content_text),
        corrected_word_count=_count_non_ws(essay.corrected_text),
    )


def _essay_to_out_bulk(essays, db: Session) -> list:
    """批量转换作文为输出结构，避免列表页 N+1 查询与逐行目录扫描。"""
    if not essays:
        return []
    user_ids = set()
    task_ids = set()
    course_ids = set()
    dirs = {}
    for e in essays:
        if e.collected_by:
            user_ids.add(e.collected_by)
        if e.reviewer_id:
            user_ids.add(e.reviewer_id)
        if e.task_id:
            task_ids.add(e.task_id)
        if e.course_id:
            course_ids.add(e.course_id)
        if e.content_file:
            fp = os.path.join(get_upload_dir(), e.content_file)
            dirs[e.id] = (os.path.dirname(fp), os.path.basename(fp))

    users = {}
    if user_ids:
        for u in db.query(User).filter(User.id.in_(user_ids)).all():
            users[u.id] = u
    tasks = {}
    if task_ids:
        from sqlalchemy.orm import joinedload
        for t in db.query(EssayTask).options(joinedload(EssayTask.course)).filter(EssayTask.id.in_(task_ids)).all():
            tasks[t.id] = t
    courses = {}
    if course_ids:
        for c in db.query(Course).filter(Course.id.in_(course_ids)).all():
            courses[c.id] = c

    dir_cache = {}
    def _corr_exists(d):
        if d not in dir_cache:
            dir_cache[d] = has_correction(d, "")
        return dir_cache[d]

    upgraded = False
    result = []
    for e in essays:
        collector = users.get(e.collected_by)
        reviewer = users.get(e.reviewer_id)
        task = tasks.get(e.task_id)
        course = courses.get(e.course_id)

        corr_exists = False
        file_path = ""
        if e.content_file:
            d, base = dirs[e.id]
            file_path = os.path.join(get_upload_dir(), e.content_file)
            corr_exists = _corr_exists(d)

        if corr_exists and e.status in ("pending", "rework") and e.content_text and e.content_text.strip():
            e.status = "confirming"
            upgraded = True

        result.append(EssayOut(
            id=e.id,
            task_id=e.task_id,
            task_name=task.name if task else "",
            course_id=e.course_id,
            course_name=course.name if course else (task.course_name if task else ""),
            grade=e.grade or "",
            essay_number=e.essay_number or 0,
            essay_title=e.essay_title or "",
            student_name=e.student_name,
            is_supplement=e.is_supplement or False,
            teaching_mode=e.teaching_mode or "线下",
            remark=e.remark or "",
            collector_note=e.collector_note or "",
            reviewer_note=e.reviewer_note or "",
            content_text=e.content_text or "",
            corrected_text=e.corrected_text or "",
            content_file=e.content_file or "",
            file_type=e.file_type or "text",
            collected_by=e.collected_by,
            collector_name=collector.nickname or collector.username if collector else "未知",
            status=e.status or "pending",
            reviewer_id=e.reviewer_id,
            reviewer_name=reviewer.nickname or reviewer.username if reviewer else "",
            corrected_at=e.corrected_at,
            created_at=e.created_at,
            file_path=file_path,
            has_correction=corr_exists,
            file_saved=e.file_saved if e.file_saved is not None else True,
            word_count=_count_non_ws(e.content_text),
            corrected_word_count=_count_non_ws(e.corrected_text),
        ))
    if upgraded:
        db.commit()
    return result


def migrate_essay_dirs_with_title(db):
    """把历史作文文件迁移到含「标题」层级的目录（修复同名多作文图片串台）。幂等，可重复执行。"""
    from ..utils.file_utils import get_essay_dir, move_content_file

    now = datetime.now()
    last_id = 0
    moved = 0
    while True:
        batch = (
            db.query(Essay)
            .filter(Essay.id > last_id, Essay.content_file.isnot(None), Essay.content_file != "")
            .order_by(Essay.id.asc())
            .limit(500)
            .all()
        )
        if not batch:
            break
        for e in batch:
            last_id = e.id
            old_dir = os.path.dirname(os.path.join(get_upload_dir(), e.content_file))
            task_name = ""
            task_created_at = None
            if e.task_id:
                task = db.query(EssayTask).filter(EssayTask.id == e.task_id).first()
                if task:
                    task_name = task.name
                    task_created_at = task.created_at
            d = e.created_at or now
            new_dir = get_essay_dir(
                str(d.year), f"{d.month}月", str(d.day),
                e.grade or "未定年级", e.essay_number, "", e.student_name, e.teaching_mode or "",
                task_name=task_name, task_created_at=task_created_at,
                essay_title=e.essay_title, is_supplement=e.is_supplement,
            )
            if os.path.abspath(old_dir) == os.path.abspath(new_dir):
                continue
            new_content_file = move_content_file(e, old_dir, new_dir, filenames=_essay_owned_filenames(e, db))
            if new_content_file:
                e.content_file = new_content_file
                moved += 1
        db.commit()
    return moved
